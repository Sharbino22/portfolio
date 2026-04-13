"""
Generate Role_A_Clinical_Lead.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Style setup ───────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)

def add_image(path, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

def bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def bold_inline(para, bold_text, normal_text):
    run = para.add_run(bold_text)
    run.bold = True
    para.add_run(normal_text)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Role A: Clinical Lead')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('I2DB Datathon 2025 — Diabetes Risk Prediction')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Reference document for the team member leading feature engineering,\n'
    'clinical narrative, and domain interpretation.'
)
run.font.size = Pt(13)
run.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Your responsibilities:')
run.font.size = Pt(12)
run.bold = True

bullets = [
    'Explain every feature engineering decision with clinical reasoning',
    'Own the "clinical story" arc of the presentation',
    'Handle clinician-facing Q&A (confounding, feature rationale, clinical workflow)',
    'Connect model outputs to real-world diabetes care',
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: DATASET OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Dataset Overview', level=1)

doc.add_paragraph(
    'We worked with electronic health record (EHR) data for 62,425 diabetes patients. Each patient '
    'has 12 months of prior data (labs, medications, visits, demographics) and a binary outcome: '
    'did their A1c become uncontrolled in the following year?'
)

add_table(
    ['Property', 'Value'],
    [
        ['Total patients', '62,425'],
        ['Features', '41 columns across 9 clinical domains'],
        ['Target', 'a1c 2025 Uncontrolled (True/False)'],
        ['Class balance', '89.4% Controlled (55,793) / 10.6% Uncontrolled (6,632)'],
        ['Imbalance ratio', '~1:8 (for every 1 uncontrolled, there are ~8 controlled)'],
    ]
)

doc.add_paragraph()
bold_para('The 9 clinical domains:')

add_table(
    ['Domain', 'Columns', 'Missingness', 'Notes'],
    [
        ['Demographics', '4', 'Near-complete', 'Birth year, sex, race, ethnicity'],
        ['A1c Labs', '11', 'A1c 1: 0%, A1c 2: 58%, A1c 3: 88%, A1c 4-5: 97%', 'Up to 5 readings + timing + 1 leakage column'],
        ['Medications', '6', '0% (zeros = no orders)', 'Metformin, insulin, GLP-1, SGLT2, sulfonylurea, DPP4'],
        ['Utilization', '3', '0%', 'ED visits, PCP visits, admissions'],
        ['Comorbidities', '2', '0%', 'CAD count, COPD count'],
        ['Cholesterol', '6', '24%', 'LDL, HDL, total + timing'],
        ['Weight/Height', '6', '73-76%', 'Values, dates, units. All kg/cm.'],
        ['Insurance', '1', '67%', '12 raw categories'],
        ['ADI', '2', '37-48%', 'Area Deprivation Index (state + national). Stored as strings.'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: THE LEAKAGE FINDING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. The Data Leakage Finding (Step 4)', level=1)

doc.add_paragraph(
    'This is the single most important data quality issue we discovered, and you need to be able '
    'to explain it clearly to judges.'
)

bold_para('What we found:')
doc.add_paragraph(
    'The column "a1c 2025-collection date-time-days from reference" records when the 2025 A1c test '
    'was collected. This column is 0% missing for uncontrolled patients but 47.7% missing for '
    'controlled patients. In other words:'
)

add_table(
    ['Column status', 'N patients', 'Controlled', 'Uncontrolled'],
    [
        ['Has a value', '35,808', '29,176', '6,632 (all uncontrolled patients)'],
        ['Missing', '26,617', '26,617 (100% controlled)', '0'],
    ]
)

doc.add_paragraph(
    'If this column is missing, the patient is always controlled. A model could achieve near-perfect '
    'accuracy simply by checking: "Is this value null? → Yes → Controlled."'
)

bold_para('Why this is leakage:')
doc.add_paragraph(
    'This column contains information about the future outcome (the 2025 A1c test). In a real-world '
    'deployment, you would be making predictions before the 2025 test happens — so you would not '
    'have this column available. A model trained with this feature would appear excellent in '
    'development but fail completely in production because the "cheat code" would not exist for new '
    'patients.'
)

bold_para('What we did:')
doc.add_paragraph(
    'We dropped this column as the very first step in data cleaning (clean_data.py, line 18). It '
    'was never used in any model.'
)

bold_para('How to explain this to judges:')
doc.add_paragraph(
    '"We identified a column that leaked future information into the training data. One hundred '
    'percent of patients missing this value were controlled — the model could have exploited this '
    'to cheat. We dropped it immediately. This is the kind of data quality issue that separates '
    'a model that works in a competition from one that could actually be deployed in a clinic."'
)

add_image('plots/step4_missingness_by_outcome.png', 5.5)
add_caption('Figure 1: Missingness rates by outcome. Note the a1c 2025 (date) row — 0% missing for uncontrolled, 47.7% for controlled.')

doc.add_paragraph(
    'Beyond the leakage column, we found that missingness itself is informative. A1c 2 is 18.8 '
    'percentage points less missing in the uncontrolled group (40.9% vs 59.7%), and A1c 3 is 12.8pp '
    'less missing. This makes clinical sense: patients with poorly controlled diabetes get more '
    'frequent A1c monitoring. We captured this signal with the n_a1c_tests feature.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: CLINICAL PATTERNS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Key Clinical Patterns from the Data', level=1)

# --- A1c risk curve ---
doc.add_heading('3.1 The A1c Risk Curve (Step 2)', level=2)

doc.add_paragraph(
    'The most important clinical finding from our exploration: risk of becoming uncontrolled rises '
    'steeply with current A1c, but plateaus above A1c 9.'
)

add_table(
    ['Current A1c', 'Risk of becoming uncontrolled', 'N patients', 'Clinical interpretation'],
    [
        ['< 5.7', '1.0%', '10,469', 'Normal range — very low risk'],
        ['5.7 – 6.5', '2.3%', '18,684', 'Pre-diabetes / well-controlled — low risk'],
        ['6.5 – 7.0', '6.5%', '9,676', 'At target — some risk begins'],
        ['7.0 – 8.0', '15.7%', '11,513', 'Above target — 1 in 6 will lose control'],
        ['8.0 – 9.0', '27.8%', '5,078', 'Poorly controlled — more than 1 in 4'],
        ['9.0 – 10.0', '32.1%', '2,591', 'Very poorly controlled — risk plateaus'],
        ['> 10.0', '32.0%', '4,414', 'Severely elevated — same plateau'],
    ]
)

add_image('plots/step2_uncontrolled_rate_by_a1c_bin.png', 5.5)
add_caption('Figure 2: Risk of becoming uncontrolled by current A1c level. Clear dose-response with plateau above 9.')

bold_para('The plateau is clinically significant.')
doc.add_paragraph(
    'Once a patient is past A1c 9, higher A1c does not mean higher future risk — the risk flattens '
    'at ~32%. This suggests that above a certain severity, factors other than the A1c number '
    'itself drive outcomes: medication adherence, engagement with care, social determinants, '
    'comorbidities. This is why we need a multivariate model — A1c alone cannot distinguish a '
    'patient at A1c 10 who will regain control from one who will not.'
)

bold_para('The 65% problem:')
doc.add_paragraph(
    '65% of patients who become uncontrolled currently have A1c below 9. If you only intervened '
    'on patients with A1c ≥ 9, you would miss nearly two-thirds of future uncontrolled patients. '
    'The model catches patients across the entire A1c spectrum by incorporating trajectory, '
    'treatment patterns, and demographics.'
)

# --- Demographics ---
doc.add_heading('3.2 Demographic Patterns (Step 3)', level=2)

add_image('plots/step3_demographics.png', 5.5)
add_caption('Figure 3: Uncontrolled rates by age, gender, and race.')

bold_para('Age:')
doc.add_paragraph(
    'Younger patients are harder to control. The under-30 group has a 14.3% uncontrolled rate '
    'versus 8.8% for those over 80. Clinically, younger diabetics often have Type 1 or aggressive '
    'early-onset Type 2, and may have worse adherence. Older patients tend to be on stable, '
    'long-term regimens. The model confirmed this: in SHAP analysis, younger age pushes predictions '
    'toward uncontrolled (SHAP = 0.093, rank #5).'
)

bold_para('Gender:')
doc.add_paragraph(
    'Males 11.8% uncontrolled versus females 9.6% — a 2.2 percentage-point gap consistent with '
    'literature showing men are slightly harder to manage (less likely to attend follow-ups, '
    'different metabolic factors). SHAP confirmed: male sex modestly pushes toward uncontrolled '
    '(SHAP = 0.055, rank #9).'
)

bold_para('Race:')
doc.add_paragraph(
    'Pacific Islander (15.4%) and American Indian/Alaska Native (14.9%) have the highest '
    'uncontrolled rates. Black or African American (11.9%) is above White (10.1%). These '
    'disparities are well-documented in diabetes literature and reflect both biological and social '
    'determinants. Important context for the equity analysis in the presentation: the model '
    'actually performs better (higher AUC, higher recall) for Black patients than White patients, '
    'meaning it does not amplify these disparities.'
)

bold_para('Hispanic ethnicity:')
doc.add_paragraph(
    'Hispanic patients: 12.8% versus non-Hispanic 10.6%. A modest but real disparity that connects '
    'to known diabetes epidemiology.'
)

# --- Medications confounding ---
doc.add_heading('3.3 The Medication Confounding Story (Step 3)', level=2)

doc.add_paragraph(
    'This is the most nuanced clinical finding and the one most likely to come up in Q&A. You need '
    'to own this explanation.'
)

bold_para('The paradox:')
doc.add_paragraph(
    'Uncontrolled patients use more medications across the board:'
)

add_table(
    ['Medication', 'Controlled patients (% with orders)', 'Uncontrolled patients (% with orders)', 'Difference'],
    [
        ['Insulin', '32.8%', '44.7%', '+11.9pp'],
        ['Metformin', '23.8%', '34.3%', '+10.5pp'],
        ['Sulfonylurea', '6.5%', '17.2%', '+10.7pp'],
        ['SGLT2', '6.6%', '11.9%', '+5.2pp'],
        ['GLP-1', '6.4%', '10.0%', '+3.7pp'],
        ['DPP4', '1.7%', '3.1%', '+1.4pp'],
    ]
)

add_image('plots/step3_comorbidities_meds.png', 5.5)
add_caption('Figure 4: Medication use and comorbidities by outcome. Note higher medication use in the uncontrolled group.')

bold_para('Why this happens — reverse causation:')
doc.add_paragraph(
    'The medications did not cause the uncontrolled A1c. The uncontrolled A1c caused the medications '
    'to be prescribed. The causal chain is:'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Patient has high A1c  →  Doctor prescribes more meds  →  Patient still has high A1c')
run.italic = True

doc.add_paragraph(
    'It is the same reason ICU patients are on more IV drips than floor patients — the drips '
    'did not make them sicker; being sicker got them the drips.'
)

bold_para('We verified this with stratified analysis:')
doc.add_paragraph(
    'Among patients with similar current A1c (e.g., A1c 6.5–8), those on more medications still '
    'have higher uncontrolled rates: 0 meds = 9.2%, 1 med = 11.1%, 2 meds = 14.5%, 3+ meds = '
    '19.4%. This is not because meds cause poor control — it is because a patient at A1c 7.5 on 3 '
    'medications has aggressive, treatment-resistant disease (they needed 3 drug classes just to get '
    'down to 7.5), while a patient at A1c 7.5 on 0 medications has mild, diet-controlled disease.'
)

bold_para('How we addressed it in feature engineering:')
doc.add_paragraph(
    'Rather than using raw medication counts (which are confounded), we created interaction features '
    'that combine A1c level with medication use:'
)

add_table(
    ['Feature', 'Definition', 'Prevalence', 'Correlation with target', 'What it captures'],
    [
        ['treatment_resistant', 'A1c ≥ 8 AND 2+ med classes', '6.4% (3,988 patients)', '+0.195 (strong)', 'Patients failing despite pharmacologic treatment — the hardest to control'],
        ['undertreated', 'A1c ≥ 8 AND 0 med classes', '4.9% (3,051 patients)', '+0.128 (moderate)', 'Patients who should be on meds but are not — possible non-adherence, access barriers, or clinical inertia'],
        ['no_medication', '0 med classes', '42.3%', '-0.118 (moderate, protective)', 'Mild, diet-controlled patients with inherently lower risk'],
    ]
)

doc.add_paragraph(
    'treatment_resistant (+0.195) became one of the strongest non-A1c predictors in the entire '
    'model, confirming that this confounding-aware feature captures a clinically real signal that '
    'raw medication counts cannot.'
)

# --- Weak signals ---
doc.add_heading('3.4 Weak Signals: Cholesterol, Comorbidities, ADI (Step 3 & 5)', level=2)

doc.add_paragraph(
    'Several feature groups showed near-zero differences between controlled and uncontrolled '
    'patients. You should know these so you are not caught off guard when a judge asks '
    '"Why didn\'t cholesterol matter?"'
)

add_table(
    ['Feature', 'Controlled mean', 'Uncontrolled mean', 'Difference', 'Why it is weak'],
    [
        ['LDL', '87.9', '89.8', '+1.8', 'LDL is managed by statins, separate from glycemic control. Well-treated in both groups.'],
        ['HDL', '112.9', '118.9', '+6.0', 'Same story — lipid management is independent of A1c.'],
        ['Total cholesterol', '163.8', '168.3', '+4.5', 'Cholesterol is LDL + HDL + triglycerides. Differences are tiny.'],
        ['ADI national rank', '64.9', '66.0', '+1.1', 'Neighborhood deprivation does not strongly predict A1c control at the individual level in this dataset.'],
        ['CAD count', '~10.4% → ~12%', '—', 'Flat', 'Having coronary artery disease does not meaningfully change glycemic risk.'],
        ['COPD count', '~10% → ~10%', '—', 'Flat', 'Same — COPD is a comorbidity but not a driver of A1c.'],
    ]
)

doc.add_paragraph(
    'These features are individually weak, but tree-based models can still extract marginal value '
    'from them in combination with stronger features. In the final model, ADI national rank had '
    'mean |SHAP| of 0.043 (rank #11) — not zero, but far below the A1c features.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: FEATURE ENGINEERING (the big one)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Feature Engineering — Complete Reference (Step 7)', level=1)

doc.add_paragraph(
    'You are responsible for explaining every feature in the model. We engineered 20 new features '
    'from the cleaned data, then combined them with 28 cleaned original columns for a total of 48 '
    'model features. Below is the complete catalog grouped by clinical domain.'
)

# --- A1c derived ---
doc.add_heading('4.1 A1c-Derived Features (8 features)', level=2)

doc.add_paragraph(
    'These are the most important features by a wide margin. Rather than feeding 5 raw A1c columns '
    '(most of which are 58-97% missing) into the model, we summarized them into clinically '
    'meaningful metrics.'
)

add_table(
    ['Feature', 'How calculated', 'Clinical reasoning', 'Corr. with target'],
    [
        ['a1c_latest', 'A1c 1 (most recent value, 0% missing)', 'Current glycemic status — the number the clinician sees first on the chart.', '+0.314'],
        ['a1c_mean', 'Mean of all available A1c readings', 'Overall glycemic burden over the observation period. Smooths noise from a single test. The #1 feature by SHAP (0.728).', '+0.333'],
        ['a1c_max', 'Maximum across all A1c readings', 'Was the patient ever severely uncontrolled? A past episode of A1c 14 is informative even if current A1c is 7.', '+0.351'],
        ['a1c_change', 'Most recent A1c − oldest A1c (2+ tests only)', 'Trajectory. Positive = worsening, negative = improving. Mean +0.20 (slight worsening on average). N=26,399.', '-0.026'],
        ['a1c_variability', 'Standard deviation of available A1c values (2+ tests)', 'Glycemic instability. A patient who swings 6→10→7 is different from one stable at 8. N=26,399.', '+0.212'],
        ['n_a1c_tests', 'Count of non-missing A1c values (1–5)', 'Monitoring intensity. 58% had 1 test, 30% had 2, 9% had 3, 3% had 4–5. More tests = closer monitoring = higher-risk patient.', '+0.145'],
        ['a1c_above_9', '1 if a1c_latest ≥ 9, else 0', 'Binary threshold for severely uncontrolled diabetes. #1 by XGBoost built-in importance (gain = 0.273).', '+0.253'],
        ['age', '2025 − birth year', 'Younger patients have higher risk (14.3% for <30 vs 8.8% for 80+). Range: 19–105, mean 65.5.', '-0.043'],
    ]
)

# --- Medication features ---
doc.add_heading('4.2 Medication Features (4 features)', level=2)

add_table(
    ['Feature', 'How calculated', 'Clinical reasoning', 'Corr. with target'],
    [
        ['total_med_classes', 'Count of med types with orders > 0 (0–6)', 'Treatment intensity. 42% on 0, 39% on 1, 14% on 2, 5% on 3+.', '+0.152'],
        ['total_med_orders', 'Sum of all medication order counts', 'Total prescription volume — captures refills and dose escalation, not just number of drugs.', '+0.100'],
        ['on_insulin', '1 if insulin orders > 0', '34% of patients. Insulin = more advanced disease requiring injectable therapy.', '+0.078'],
        ['on_newer_drugs', '1 if GLP-1 or SGLT2 orders > 0', '13% of patients. Newer, expensive agents often reserved for harder-to-treat cases.', '+0.071'],
    ]
)

# --- Treatment patterns ---
doc.add_heading('4.3 Treatment Pattern Features (3 features)', level=2)

doc.add_paragraph(
    'These are the features that address the medication confounding directly. They combine A1c '
    'level with medication count to create clinically meaningful categories.'
)

add_table(
    ['Feature', 'How calculated', 'Clinical reasoning', 'Corr. with target'],
    [
        ['treatment_resistant', 'A1c ≥ 8 AND total_med_classes ≥ 2', 'The patient on multiple meds who still cannot control their A1c. 6.4% of patients. One of the strongest non-A1c features. XGBoost importance rank #4.', '+0.195'],
        ['undertreated', 'A1c ≥ 8 AND total_med_classes = 0', 'The patient who should be on meds but is not. May indicate non-adherence, access barriers, or clinical inertia. 4.9% of patients.', '+0.128'],
        ['no_medication', 'total_med_classes = 0', 'Mild, diet-controlled patients. 42.3%. These patients have inherently lower risk — being on zero meds is a sign of mild disease, not under-treatment.', '-0.118'],
    ]
)

# --- Body metrics ---
doc.add_heading('4.4 Body Metrics (2 features)', level=2)

add_table(
    ['Feature', 'How calculated', 'Clinical reasoning', 'Corr. with target'],
    [
        ['bmi', 'Weight (kg) / (height in m)². Capped at 10–80.', 'BMI mean 32.6 (obese range). Only 23.5% of patients have both weight+height. Despite high missingness, SHAP rank #6 (0.081).', '-0.002'],
        ['has_bmi', '1 if BMI can be calculated', 'Whether vitals were documented. Patients seen more frequently may be more likely to have measurements.', '-0.002'],
    ]
)

# --- Utilization ---
doc.add_heading('4.5 Utilization Features (3 features)', level=2)

add_table(
    ['Feature', 'How calculated', 'Clinical reasoning', 'Corr. with target'],
    [
        ['total_encounters', 'ED visits + PCP visits + admissions', 'Combines three correlated utilization columns (PCP and admissions r=0.73).', '+0.020'],
        ['any_admission', '1 if admissions > 0', '32% hospitalized. Admission is a marker of disease severity.', '-0.000'],
        ['any_comorbidity', '1 if CAD > 0 or COPD > 0', '23% have at least one comorbidity. Captures overall disease burden.', '+0.012'],
    ]
)

# --- Kept from original ---
doc.add_heading('4.6 Original Features Kept (28 columns)', level=2)

doc.add_paragraph(
    'In addition to the 20 engineered features, 28 columns from the cleaned dataset were kept as-is:'
)

add_table(
    ['Group', 'Features', 'Count'],
    [
        ['Raw medications', 'glp-1, insulin, metformin, sglt2, sulfonylurea, dpp4 order counts', '6'],
        ['Raw utilization', 'ED visit count, PCP visit count, admission count', '3'],
        ['Comorbidities', 'cad-count, copd-count', '2'],
        ['Cholesterol', 'LDL, HDL, total cholesterol', '3'],
        ['ADI', 'ADI national rank', '1'],
        ['Demographics (encoded)', 'is_male, is_hispanic', '2'],
        ['Race (one-hot)', '7 race indicator columns', '7'],
        ['Insurance (one-hot)', 'Managed Care, Medicaid, Medicare, Other', '4'],
    ]
)

doc.add_paragraph(
    'Total: 20 engineered + 28 original = 48 features in the final model.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: SHAP AND CLINICAL INTERPRETATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. What the Model Learned (SHAP Analysis)', level=1)

doc.add_paragraph(
    'SHAP (SHapley Additive exPlanations) shows exactly how each feature pushes each patient\'s '
    'prediction toward "uncontrolled" or "controlled." This is your key slide for the presentation — '
    'it proves the model reasons like a clinician.'
)

add_image('plots/step13_shap_summary.png', 5.5)
add_caption('Figure 5: SHAP beeswarm plot. Each dot = one patient. X-axis = push toward uncontrolled (right) or controlled (left). Color = feature value (red = high, blue = low).')

bold_para('How to read this plot (row by row):')

add_table(
    ['Feature', 'SHAP rank', 'Pattern', 'Plain English'],
    [
        ['a1c_mean', '#1 (0.728)', 'Red right, blue left', 'Higher average A1c = strong push toward uncontrolled. This is the single most important factor.'],
        ['a1c_max', '#2 (0.501)', 'Red right, blue left', 'The worst A1c reading matters — a bad history increases risk even if current A1c has improved.'],
        ['a1c_latest', '#3 (0.218)', 'Red right, blue left', 'Most recent A1c. Adds recency information beyond the average.'],
        ['a1c_change', '#4 (0.108)', 'Red right (A1c went UP)', 'Worsening trajectory strongly predicts loss of control. Improving trajectory is protective.'],
        ['age', '#5 (0.093)', 'Blue right (YOUNG)', 'Younger patients are higher risk — counterintuitive but consistent with aggressive early-onset disease.'],
        ['bmi', '#6 (0.081)', 'Red right (high BMI)', 'Higher BMI pushes toward uncontrolled. Works despite 77% missing data.'],
        ['a1c_variability', '#7 (0.078)', 'Red right (unstable)', 'Swinging A1c = higher risk. Glycemic stability matters.'],
        ['a1c_above_9', '#8 (0.057)', 'Red right', 'The binary "severely uncontrolled" flag.'],
        ['is_male', '#9 (0.055)', 'Red right (male)', 'Being male modestly increases predicted risk.'],
        ['sulfonylurea', '#10 (0.047)', 'Red right', 'Sulfonylurea use pushes toward uncontrolled — confounded (prescribed to sicker patients).'],
    ]
)

bold_para('The clinical story for judges:')
doc.add_paragraph(
    '"The model\'s decision-making mirrors clinical reasoning. It first evaluates A1c history — '
    'the average, the worst episode, the most recent reading, and the trajectory. It then considers '
    'whether the patient is responding to treatment (treatment_resistant), considers demographics '
    '(age, sex), and factors in social context (ADI, insurance). This is not a black box — it '
    'reasons the way a clinician would, just faster and at scale."'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6: PRESENTATION TALKING POINTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Your Presentation Talking Points', level=1)

doc.add_paragraph(
    'As Clinical Lead, you will likely present or co-present slides 2 (the problem), 4 (A1c risk '
    'curve), 5 (feature engineering), 9 (SHAP), and 10 (equity). Below are ready-to-use talking '
    'points for each.'
)

doc.add_heading('Slide 2: The Clinical Problem', level=2)
doc.add_paragraph(
    '"37 million Americans have diabetes, and glycemic control is the cornerstone of management. '
    'When A1c becomes uncontrolled, the risk of complications rises sharply — kidney disease, '
    'retinopathy, neuropathy, cardiovascular events. Currently, we discover loss of control '
    'reactively: the patient comes in, labs are drawn, and the A1c comes back elevated. What if we '
    'could identify these patients before it happens? That is what our model does — it flags '
    'patients who are likely to lose control in the next year, enabling proactive intervention."'
)

doc.add_heading('Slide 4: A1c is Necessary but Not Sufficient', level=2)
doc.add_paragraph(
    '"Current A1c is the strongest single predictor, with a clear dose-response from 1% risk at '
    'normal levels to 32% at A1c above 9. But there are two key insights. First, risk plateaus '
    'above A1c 9 — meaning at that severity, other factors like medication response and engagement '
    'matter more than the number itself. Second, 65% of patients who become uncontrolled currently '
    'have A1c below 9. If you only intervened on the highest-A1c patients, you would miss '
    'two-thirds of the ones who need help. That is why we built a multivariate model."'
)

doc.add_heading('Slide 5: Feature Engineering', level=2)
doc.add_paragraph(
    '"We engineered 20 features guided by clinical reasoning. The most impactful was '
    'treatment_resistant — defined as A1c above 8 despite being on two or more medication classes. '
    'Raw medication counts are confounded: sicker patients get more meds. But when you combine A1c '
    'level with medication count, you capture a genuinely predictive signal. A patient at A1c 7.5 '
    'on 3 medications has aggressive disease that required 3 drug classes just to reach 7.5. A '
    'patient at A1c 7.5 on zero medications has mild disease managed by diet. Same A1c, completely '
    'different risk profile. Treatment_resistant became one of the strongest non-A1c features in '
    'our model."'
)

doc.add_heading('Slide 9: What the Model Learned', level=2)
doc.add_paragraph(
    '"SHAP analysis shows the model reasons like a clinician. It first checks A1c history — the '
    'average, the worst reading, the trend. Then it asks: is this patient responding to treatment? '
    'Then demographics and social context. The top 7 features are all clinically interpretable. '
    'This is not a black box — every prediction can be traced to specific clinical factors."'
)

doc.add_heading('Slide 10: Equity', level=2)
doc.add_paragraph(
    '"We tested the model across race, gender, age, insurance, and neighborhood deprivation. AUC '
    'ranges from 0.83 to 0.90 across all subgroups — no group is left behind. Notably, the model '
    'has higher sensitivity for Black patients and Medicaid patients — the populations facing the '
    'greatest diabetes disparities. This means the tool could help close equity gaps rather than '
    'widen them."'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7: Q&A PREP
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Anticipated Q&A — Your Answers', level=1)

doc.add_paragraph(
    'These are the questions most likely to be directed at the Clinical Lead. Rehearse these answers.'
)

questions = [
    (
        'Q: Why not just use A1c as the predictor? Why do you need a model?',
        'A: A1c is the strongest single feature, but 65% of patients who become uncontrolled '
        'currently have A1c below 9. A threshold-based rule would miss most of them. The model '
        'combines A1c with trajectory (is it getting worse?), treatment response (are meds working?), '
        'monitoring intensity (how closely is the patient being watched?), and demographics to '
        'catch patients that A1c alone would miss.'
    ),
    (
        'Q: How do you handle the medication confounding?',
        'A: We do not use raw medication counts as simple predictors. Instead, we engineered '
        'interaction features that combine A1c level with medication use. Treatment_resistant '
        '(A1c ≥ 8 despite 2+ drug classes) captures patients failing pharmacologic treatment. '
        'Undertreated (A1c ≥ 8 with zero drugs) captures those who may not be adherent or are '
        'facing access barriers. These features separate "sick on meds" from "mild off meds" — '
        'the distinction the raw counts cannot make.'
    ),
    (
        'Q: Why is precision only 27%? Is the model bad?',
        'A: No — 27% precision is expected for screening with a 10.6% base rate. For comparison, '
        'mammography flags ~10% of women but only ~0.5% have cancer — that is about 5% precision. '
        'Our 27% is actually strong. The clinical cost of a false alarm (a 15-minute chart review) '
        'is far less than missing a patient who loses glycemic control. And the NPV is 97% — when '
        'the model says a patient is safe, it is right 97% of the time.'
    ),
    (
        'Q: What about cholesterol and comorbidities? Did they help?',
        'A: Individually, cholesterol and comorbidities are weak predictors — mean values are '
        'nearly identical between groups. Cholesterol is managed by statins independently of A1c. '
        'CAD and COPD are comorbidities but not drivers of glycemic control. However, the tree-based '
        'model can still extract marginal value from these in combination with stronger features. '
        'ADI national rank, for example, had mean |SHAP| of 0.043 — not zero, but far below A1c.'
    ),
    (
        'Q: Is the model biased against any demographic group?',
        'A: No. We tested across race, gender, age, insurance, and ADI. AUC ranges from 0.83 to '
        '0.90. The model actually has higher sensitivity for Black patients (85% recall) than White '
        '(79%), and highest recall for Medicaid patients (91–93%). If anything, it catches more '
        'at-risk patients in underserved populations.'
    ),
    (
        'Q: How would this work in a real clinic?',
        'A: Monthly batch scoring on the diabetes panel. The model outputs a risk probability '
        'for each patient. Those above the threshold get flagged. A nurse reviews the chart, '
        'confirms the risk, and the provider proactively adjusts the treatment plan — medication '
        'escalation, closer follow-up, care coordination. The threshold is adjustable: lower for '
        'population screening (catch everyone), higher for targeted intervention (fewer flags).'
    ),
    (
        'Q: You dropped the leakage column. How did you find it?',
        'A: During missingness analysis (Step 4), we compared missing rates between controlled and '
        'uncontrolled groups. The a1c 2025 collection date column was 0% missing for uncontrolled '
        'patients but 47.7% missing for controlled patients. That is a perfect signal: if this '
        'value is missing, the patient is always controlled. We confirmed with a cross-tabulation '
        'and dropped it immediately. This is why careful data exploration matters — without this '
        'step, the model would have learned a shortcut that would not work in production.'
    ),
    (
        'Q: What does treatment_resistant really capture clinically?',
        'A: It identifies patients with refractory diabetes — those who are on multiple medication '
        'classes but still cannot achieve glycemic control. These are the patients who need the most '
        'clinical attention: medication optimization, specialist referral, adherence support, or '
        'investigation of complicating factors (stress, diet, other medications like steroids). '
        'The model correctly identifies them as highest risk beyond what A1c alone would suggest.'
    ),
]

for q, a in questions:
    bold_para(q)
    doc.add_paragraph(a)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# QUICK REFERENCE CARD
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('Quick Reference Card', level=1)

doc.add_paragraph('Pin this page. These are the numbers you need at your fingertips.')

bold_para('Key numbers:')
add_table(
    ['Metric', 'Value'],
    [
        ['Patients', '62,425'],
        ['Features (final model)', '48'],
        ['Engineered features', '20'],
        ['Model', 'XGBoost (tuned)'],
        ['AUC-ROC', '0.852'],
        ['Sensitivity (Recall)', '80.6%'],
        ['NPV', '97.0%'],
        ['Precision', '27.0%'],
        ['Top predictor', 'a1c_mean (SHAP = 0.728)'],
        ['Best non-A1c predictor', 'treatment_resistant (SHAP = ~0.05, built-in rank #4)'],
        ['Class balance', '10.6% uncontrolled'],
        ['Leakage column', 'Dropped (a1c 2025 collection date)'],
    ]
)

doc.add_paragraph()
bold_para('One-liner for each feature group:')
add_table(
    ['Group', 'Signal strength', 'One-liner'],
    [
        ['A1c features', 'Dominant', 'Mean, max, latest, trajectory, variability — these carry 70%+ of predictive power'],
        ['Medications', 'Strong (confounded)', 'More meds = sicker patient, not worse treatment. Captured by treatment_resistant.'],
        ['Demographics', 'Moderate', 'Younger, male, minority → modestly higher risk'],
        ['ED visits', 'Moderate', 'Dose-response: more ED visits = higher risk (marker of instability)'],
        ['Cholesterol', 'Weak', 'Managed by statins independently of A1c. Nearly identical between groups.'],
        ['Comorbidities', 'Weak', 'CAD/COPD are present but not drivers of glycemic control'],
        ['ADI', 'Weak individually', 'Neighborhood deprivation adds marginal signal in combination with other features'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save('Role_A_Clinical_Lead.docx')
print('Saved: Role_A_Clinical_Lead.docx')
