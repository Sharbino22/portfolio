"""
Generate Role_C_Data_Explorer_Presentation.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)

def add_image(path, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_thumb(path, width=3.5):
    """Smaller thumbnail version for gallery views."""
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

def bold_para(text):
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    return p

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Role C: Data Explorer & Presentation')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('I2DB Datathon 2025 — Diabetes Risk Prediction')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Reference document for the team member designing slides,\n'
    'managing the visual story, and presenting exploratory findings.'
)
run.font.size = Pt(13)
run.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Your responsibilities:')
run.font.size = Pt(12)
run.bold = True

for b in [
    'Build the slide deck using the 13-slide structure and presentation-ready plots',
    'Know every exploration finding well enough to narrate the "data story"',
    'Manage the plot inventory — which figure goes where',
    'Design clean, readable slides that communicate findings to a mixed audience',
]:
    p = doc.add_paragraph(b, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: EXPLORATION FINDINGS SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Exploration Findings Summary (Steps 1–5)', level=1)

doc.add_paragraph(
    'This section gives you the complete data story from loading to correlations. You need to know '
    'this cold — it forms the first half of the presentation narrative.'
)

# --- Step 1 ---
doc.add_heading('1.1 Step 1: Loading and Initial Inspection', level=2)

add_table(
    ['Finding', 'Detail'],
    [
        ['Shape', '62,425 patients × 41 features + 1 target column'],
        ['ID match', 'Both files share identical patient IDs in the same order'],
        ['Feature groups', '9 groups: demographics (4), comorbidities (2), A1c labs (11), weight/height (6), cholesterol (6), utilization (3), insurance (1), medications (6), ADI (2)'],
        ['Zero-missing columns', '14 columns: all medications, utilization, comorbidities, A1c 1, birth year'],
        ['Heaviest missingness', 'A1c 4-5 (97%), height (76%), weight (73%), insurance (67%), A1c 2 (58%)'],
        ['Type issues', 'ADI columns stored as strings (e.g., "10.0") — needed numeric conversion'],
        ['Leakage flag', '"a1c 2025-collection date-time-days from reference" — suspicious, investigated in Step 4'],
    ]
)

# --- Step 2 ---
doc.add_heading('1.2 Step 2: Target Variable Analysis', level=2)

bold_para('Class balance:')
doc.add_paragraph(
    '89.4% Controlled (55,793) vs 10.6% Uncontrolled (6,632). Roughly 1:8 ratio. A dummy model '
    'predicting "controlled" for everyone would be 89.4% accurate but catch 0% of at-risk patients.'
)

bold_para('A1c by outcome:')
add_table(
    ['Metric', 'Controlled', 'Uncontrolled', 'Difference'],
    [
        ['Mean A1c (most recent)', '6.92', '8.78', '+1.86 points'],
        ['Median A1c', '6.5', '8.3', '+1.8 points'],
        ['% with A1c ≥ 9', '9.0%', '35.5%', '+26.5pp'],
    ]
)

bold_para('A1c risk gradient (the most important finding):')
add_table(
    ['A1c range', 'Risk', 'N patients'],
    [
        ['< 5.7', '1.0%', '10,469'],
        ['5.7–6.5', '2.3%', '18,684'],
        ['6.5–7.0', '6.5%', '9,676'],
        ['7.0–8.0', '15.7%', '11,513'],
        ['8.0–9.0', '27.8%', '5,078'],
        ['9.0–10.0', '32.1%', '2,591'],
        ['> 10.0', '32.0%', '4,414'],
    ]
)

doc.add_paragraph(
    'Two key insights: (1) Risk plateaus above A1c 9 — other factors matter more at that severity. '
    '(2) 65% of future-uncontrolled patients currently have A1c < 9 — a threshold rule would miss most.'
)

# --- Step 3 ---
doc.add_heading('1.3 Step 3: Feature Group Exploration', level=2)

bold_para('Feature signal ranking (from strongest to weakest):')
add_table(
    ['Rank', 'Feature group', 'Signal strength', 'Key evidence'],
    [
        ['1', 'A1c values', 'Dominant', 'Mean A1c difference of 1.86 between groups. Correlations +0.31 to +0.43.'],
        ['2', 'Medications', 'Strong (confounded)', 'Insulin +12pp, metformin +11pp, sulfonylurea +11pp difference. But sicker patients get more meds.'],
        ['3', 'Demographics', 'Moderate', 'Age (<30: 14.3% vs 80+: 8.8%), male (+2.2pp), race disparities (Pacific Islander 15.4% vs White 10.1%).'],
        ['4', 'ED visits', 'Moderate', 'Dose-response: 0 visits = 9.7%, 5+ visits = 18.4%.'],
        ['5', 'Cholesterol', 'Weak', 'LDL diff +1.8, HDL diff +6.0, total chol diff +4.5. Nearly identical distributions.'],
        ['6', 'Comorbidities', 'Weak', 'CAD and COPD show flat rates across counts.'],
        ['7', 'ADI', 'Weak', 'National rank diff only +1.1 (64.9 vs 66.0).'],
    ]
)

bold_para('The medication confounding (critical to understand):')
doc.add_paragraph(
    'Uncontrolled patients are on MORE meds, not fewer. This is reverse causation — high A1c causes '
    'more prescribing, not the other way around. Verified by stratified analysis: among patients at '
    'the same A1c level, those on more meds still have worse outcomes because they have more '
    'aggressive underlying disease. This led to the treatment_resistant feature.'
)

# --- Step 4 ---
doc.add_heading('1.4 Step 4: Missingness Analysis', level=2)

bold_para('Missingness is NOT random:')
add_table(
    ['Column', 'Controlled % missing', 'Uncontrolled % missing', 'Difference', 'Interpretation'],
    [
        ['a1c 2025 date (LEAKAGE)', '47.7%', '0.0%', '-47.7pp', 'DROPPED — perfect predictor of outcome, not available in production'],
        ['A1c 2', '59.7%', '40.9%', '-18.8pp', 'Uncontrolled patients get more repeat labs'],
        ['A1c 3', '89.1%', '76.3%', '-12.8pp', 'Same pattern — more monitoring for sicker patients'],
        ['Cholesterol (HDL/total)', '25.1%', '20.0%', '-5.1pp', 'More labs drawn overall for uncontrolled'],
        ['Weight/height', '~73-76%', '~74-76%', '~0', 'Similar — missingness not informative here'],
        ['Insurance', '67.1%', '66.2%', '~0', 'Similar'],
    ]
)

doc.add_paragraph(
    'Key insight: the count of non-missing A1c values (n_a1c_tests) becomes a useful feature — more '
    'tests = closer monitoring = higher-risk patient.'
)

# --- Step 5 ---
doc.add_heading('1.5 Step 5: Correlations', level=2)

bold_para('Top correlations with target:')
add_table(
    ['Feature', 'Correlation', 'Strength'],
    [
        ['A1c 5 (oldest reading)', '+0.432', 'Strong'],
        ['A1c 4', '+0.404', 'Strong'],
        ['A1c 3', '+0.387', 'Strong'],
        ['A1c 2', '+0.386', 'Strong'],
        ['A1c 1 (most recent)', '+0.314', 'Strong'],
        ['Sulfonylurea orders', '+0.109', 'Moderate'],
        ['Insulin orders', '+0.082', 'Moderate'],
        ['Metformin orders', '+0.066', 'Moderate'],
        ['SGLT2 orders', '+0.050', 'Weak-moderate'],
        ['Birth year (younger)', '+0.043', 'Weak'],
        ['Everything else', '< 0.04', 'Weak'],
    ]
)

bold_para('Redundant feature pairs we addressed:')
add_table(
    ['Pair', 'Correlation', 'Action'],
    [
        ['ADI state vs national', 'r = 0.964', 'Dropped state, kept national'],
        ['LDL vs total cholesterol', 'r = 0.772', 'Kept both (model handles it)'],
        ['PCP visits vs admissions', 'r = 0.733', 'Created total_encounters'],
        ['A1c readings (various)', 'r = 0.52–0.64', 'Aggregated into mean, max, change, variability'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: COMPLETE PLOT INVENTORY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Complete Plot Inventory', level=1)

doc.add_paragraph(
    'Every PNG file in the plots/ folder, organized by category. Use this as your master reference '
    'when building slides.'
)

doc.add_heading('2.1 Exploration Plots (Steps 2–5)', level=2)

add_table(
    ['Filename', 'Step', 'What it shows', 'Best slide for'],
    [
        ['step2_class_balance.png', '2', 'Bar chart: 55,793 controlled vs 6,632 uncontrolled (89.4% vs 10.6%)', 'Slide 2 (backup) or appendix'],
        ['step2_a1c1_by_outcome.png', '2', 'Histogram + box plot of most recent A1c split by outcome', 'Slide 4 (backup)'],
        ['step2_uncontrolled_rate_by_a1c_bin.png', '2', 'Bar chart: risk of uncontrolled by A1c bin (<5.7 through >10)', 'Slide 4 (backup) — pres_2 version is cleaner'],
        ['step3_demographics.png', '3', '4-panel: age distribution, age bins, gender, race — all by outcome', 'Appendix or equity backup'],
        ['step3_comorbidities_meds.png', '3', '4-panel: CAD, COPD, medication use, med classes — by outcome', 'Slide 5 (feature engineering discussion)'],
        ['step3_util_chol_adi.png', '3', '6-panel: ED visits, PCP, admissions, LDL, HDL, ADI — by outcome', 'Appendix'],
        ['step4_missingness_heatmap.png', '4', 'Heatmap of missing values for 500 patients sorted by # missing', 'Appendix'],
        ['step4_missingness_by_outcome.png', '4', 'Side-by-side bars: missingness rate controlled vs uncontrolled', 'Leakage discussion (if asked)'],
        ['step5_correlation_matrix.png', '5', 'Full correlation heatmap of all numeric features + target', 'Appendix'],
        ['step5_target_correlations.png', '5', 'Horizontal bars: correlation of each feature with target, ranked', 'Appendix or methods backup'],
    ]
)

doc.add_heading('2.2 Model Plots (Steps 9–12)', level=2)

add_table(
    ['Filename', 'Step', 'What it shows', 'Best slide for'],
    [
        ['step9_baseline_roc_pr.png', '9', 'ROC + PR curves for logistic regression (6 feat vs 48 feat)', 'Appendix (superseded by step 10/11 plots)'],
        ['step10_model_comparison.png', '10', 'ROC + PR: logistic regression vs Random Forest vs XGBoost (defaults)', 'Slide 6 (backup) — pres_3 is cleaner'],
        ['step11_tuned_comparison.png', '11', 'ROC + PR: all 4 models (LR, RF, XGB default, XGB tuned)', 'Slide 6 — shows full model evolution'],
        ['step12_full_evaluation.png', '12', '4-panel: confusion matrix, ROC, PR curve, calibration', 'Slide 7/8 — comprehensive evaluation'],
        ['step12_probability_distribution.png', '12', 'Histogram of predicted probabilities by true outcome', 'Slide 8 or appendix — shows model separation'],
    ]
)

doc.add_heading('2.3 Interpretation Plots (Step 13)', level=2)

add_table(
    ['Filename', 'Step', 'What it shows', 'Best slide for'],
    [
        ['step13_feature_importance.png', '13', 'Top 20 features by XGBoost gain (built-in importance)', 'Appendix (SHAP is better for presentation)'],
        ['step13_shap_summary.png', '13', 'SHAP beeswarm: how each feature pushes predictions, all patients', 'Slide 9 (backup) — pres_5 is cleaner top-10 version'],
        ['step13_shap_bar.png', '13', 'Mean |SHAP| bar chart: average impact per feature', 'Slide 9 (alternative to beeswarm)'],
    ]
)

doc.add_heading('2.4 Fairness Plots (Step 14)', level=2)

add_table(
    ['Filename', 'Step', 'What it shows', 'Best slide for'],
    [
        ['step14_subgroup_analysis.png', '14', '6-panel: AUC by race, gender, age, ADI, insurance + recall by race', 'Slide 10 (backup) — very detailed'],
    ]
)

doc.add_heading('2.5 Presentation-Ready Plots (pres_ prefix)', level=2)

doc.add_paragraph(
    'These 7 figures were specifically designed for slides: larger fonts, cleaner layouts, '
    'publication-quality styling. Use these as your primary slide content.'
)

add_table(
    ['Filename', 'Target slide', 'What it shows'],
    [
        ['pres_1_problem.png', 'Slide 2: The Clinical Problem', 'Clean bar chart: 55,793 controlled vs 6,632 uncontrolled with counts and percentages'],
        ['pres_2_risk_by_a1c.png', 'Slide 4: A1c is the Strongest Signal', 'Color-gradient bar chart: risk by A1c bin from 1% to 32%'],
        ['pres_3_model_roc.png', 'Slide 6: Model Comparison', 'ROC curves for LR (0.828), RF (0.847), XGBoost tuned (0.852)'],
        ['pres_4_confusion.png', 'Slide 8: Understanding Predictions', 'Confusion matrix with plain-language labels (TN/FP/FN/TP)'],
        ['pres_5_shap.png', 'Slide 9: What Drives Predictions', 'SHAP beeswarm (top 10 features only, cleaner than step13 version)'],
        ['pres_6_fairness.png', 'Slide 10: The Model is Equitable', 'Side-by-side: AUC by race + AUC by gender with overall reference line'],
        ['pres_7_metrics.png', 'Slide 7: Final Model Performance', '5 bars: AUC 85.2%, Sensitivity 80.6%, NPV 97.0%, Specificity 74.2%, Precision 27.0%'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: PRESENTATION-READY FIGURE GALLERY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Presentation Figure Gallery', level=1)

doc.add_paragraph(
    'Thumbnails of all 7 presentation-ready figures. Review these to understand what is already '
    'built. You can use these directly in PowerPoint — they are 200 DPI PNGs.'
)

gallery = [
    ('pres_1_problem.png', 'Slide 2 — The Clinical Problem', 'Shows the 1:8 class imbalance. Open with this to frame the challenge.'),
    ('pres_2_risk_by_a1c.png', 'Slide 4 — A1c Risk Gradient', 'The color gradient from green to red tells the story visually. Point out the plateau above A1c 9.'),
    ('pres_3_model_roc.png', 'Slide 6 — Model Comparison', 'Three ROC curves. XGBoost (blue) pulls ahead. Keep it simple — "blue wins."'),
    ('pres_4_confusion.png', 'Slide 8 — Confusion Matrix', 'The labels say "correctly cleared," "false alarm," "missed," and "caught!" — judge-friendly language.'),
    ('pres_5_shap.png', 'Slide 9 — SHAP Analysis', 'Top 10 features only for readability. Red dots right = high value pushes toward uncontrolled.'),
    ('pres_6_fairness.png', 'Slide 10 — Equity', 'Side-by-side AUC by race and gender. Dashed line shows overall AUC.'),
    ('pres_7_metrics.png', 'Slide 7 — Key Metrics', 'Five bars color-coded: blue (AUC), green (sensitivity, NPV), orange (specificity), red (precision).'),
]

for filename, slide_label, tip in gallery:
    doc.add_heading(slide_label, level=3)
    add_thumb(f'plots/{filename}', 4.5)
    add_caption(f'{filename}')
    p = doc.add_paragraph()
    run = p.add_run('Design tip: ')
    run.bold = True
    p.add_run(tip)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: 13-SLIDE STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Full 13-Slide Presentation Structure', level=1)

doc.add_paragraph(
    'This is the complete slide map. For each slide: the title, the figure(s) to use, and the key '
    'message. Aim for 10–15 minutes total (roughly 1 minute per slide, 2 minutes for slides 7–9).'
)

slides = [
    ('Slide 1: Title', 'None',
     '"Predicting Loss of Glycemic Control: A Machine Learning Approach Using 12 Months of EHR Data." Team name, institution, I2DB Datathon 2026. Clean, minimal.'),

    ('Slide 2: The Clinical Problem', 'pres_1_problem.png',
     '62,425 patients, 10.6% become uncontrolled. Frame it as: "Currently discovered reactively — what if we could predict it?" The bar chart makes the imbalance visceral.'),

    ('Slide 3: The Data', 'None — use a clean table/bullets',
     '41 features across 9 domains. List the domains with icons if possible. Mention significant missingness (weight 73%, A1c 2 58%, insurance 67%). Keep this slide fast — 30 seconds.'),

    ('Slide 4: A1c is the Strongest Signal', 'pres_2_risk_by_a1c.png',
     'The gradient bar chart tells the story. Two takeaways: (1) plateau above A1c 9, (2) 65% of future-uncontrolled have A1c < 9 now. "A1c is necessary but not sufficient."'),

    ('Slide 5: Feature Engineering', 'None — use a table',
     'Table of top engineered features with clinical reasoning. Highlight treatment_resistant. This is where you show clinical thinking, not just data science. The confounding story lives here.'),

    ('Slide 6: Model Comparison', 'pres_3_model_roc.png',
     'Three curves: LR (gray, 0.828), RF (orange, 0.847), XGBoost tuned (blue, 0.852). "We tested three approaches; XGBoost with tuning won."'),

    ('Slide 7: Final Model Performance', 'pres_7_metrics.png',
     'Five bars. Lead with the green ones: "81% sensitivity — catches 4 out of 5. 97% NPV — when the model clears a patient, it is right 97% of the time." Address precision proactively: "27% is expected for screening at 10.6% base rate."'),

    ('Slide 8: Understanding the Predictions', 'pres_4_confusion.png',
     'Walk through each quadrant: 8,275 correctly cleared (time saved), 1,069 caught (intervention opportunities), 257 missed (area for improvement), 2,884 false alarms (cost = chart review). "Screen 3,953 to find 1,069."'),

    ('Slide 9: What Drives the Predictions?', 'pres_5_shap.png',
     'SHAP beeswarm (top 10). Walk through top 3 rows: a1c_mean, a1c_max, a1c_latest. Note age (younger = riskier). "The model reasons like a clinician: A1c history first, then treatment response, then demographics."'),

    ('Slide 10: The Model is Equitable', 'pres_6_fairness.png',
     'AUC 0.83–0.88 across subgroups. "No group is left behind. Higher sensitivity for Black and Medicaid patients — the most at-risk populations." This slide wins trust with judges.'),

    ('Slide 11: Clinical Application', 'None — use a table',
     'Threshold table: 0.15 catches 96%, 0.30 catches 91%, 0.50 catches 81%. Suggested workflow: monthly scoring → flag → RN review → proactive treatment. Make it feel deployable.'),

    ('Slide 12: Limitations and Future Work', 'None — bullets',
     'Single-institution data, missing weight (73%), small Pacific Islander/American Indian samples, model predicts risk not causation. Future: external validation, medication adherence data, CDS integration.'),

    ('Slide 13: Summary', 'None — 5 bullet points',
     '(1) 10.6% lose control annually. (2) 48 features from routine EHR data. (3) AUC 0.852, 81% recall, 97% NPV. (4) Equitable across subgroups. (5) Enables proactive care.'),
]

add_table(
    ['Slide', 'Figure', 'Key message (one sentence)'],
    [(s, f, m.split('.')[0] + '.' if '.' in m else m[:80]) for s, f, m in slides]
)

doc.add_paragraph()

for title, figure, message in slides:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    run = p.add_run('Figure: ')
    run.bold = True
    p.add_run(figure)

    p = doc.add_paragraph()
    run = p.add_run('Message: ')
    run.bold = True
    p.add_run(message)

    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: MISSINGNESS PATTERNS SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Missingness Patterns — Quick Reference', level=1)

doc.add_paragraph(
    'If asked "how did you handle missing data?" — here is the complete picture.'
)

add_table(
    ['Column group', '% Missing', 'How handled', 'Notes'],
    [
        ['A1c 1 (most recent)', '0%', 'Used directly as a1c_latest', 'Fully populated'],
        ['A1c 2', '57.7%', 'Aggregated into mean/max/change/variability', 'Missing = patient had only 1 test'],
        ['A1c 3', '87.8%', 'Same aggregation', 'Missing = patient had only 1-2 tests'],
        ['A1c 4-5', '97.2%', 'Same aggregation', 'Only 2.8% of patients had 4+ tests'],
        ['Weight', '73.3%', 'Used to compute BMI (available for 23.5%)', 'has_bmi feature captures missingness'],
        ['Height', '75.9%', 'Used to compute BMI', 'Same as weight'],
        ['Insurance', '67.0%', 'One-hot encoded (missing = all 4 dummies = 0)', 'Missing insurance is the implicit reference category'],
        ['ADI national rank', '48.1%', 'Converted to numeric; XGBoost handles NaN natively', 'No imputation needed for XGBoost'],
        ['Cholesterol (LDL/HDL/total)', '24-25%', 'Kept as-is; XGBoost handles NaN natively', 'Weak predictors regardless'],
        ['a1c 2025 date (LEAKAGE)', '42.6%', 'DROPPED entirely', '100% of missing = controlled. Data leakage.'],
        ['Medications, utilization, comorbidities', '0%', 'Used directly', 'Zeros = no orders/visits/diagnoses'],
        ['Demographics (gender, race, ethnicity)', '0-0.1%', 'Encoded; rare missing values preserved as NaN', 'Only 4 missing gender, 27 missing ethnicity'],
    ]
)

doc.add_paragraph(
    'For logistic regression (Step 9), we used median imputation via SimpleImputer. For XGBoost '
    '(Steps 10–12), we used native NaN handling — XGBoost learns which direction to send missing '
    'values at each split, which is generally superior to imputation.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6: WHAT JUDGES LOOK FOR
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. What Judges Typically Look For', level=1)

doc.add_paragraph(
    'Based on datathon judging norms and the competition description, judges typically evaluate '
    'along these dimensions. Design your slides to hit every one.'
)

add_table(
    ['Dimension', 'What they want to see', 'Where we deliver it'],
    [
        ['Clinical relevance', 'Does the problem matter? Is the approach clinically sensible?', 'Slide 2 (problem framing), Slide 5 (feature engineering with clinical reasoning), Slide 11 (clinical workflow)'],
        ['Technical rigor', 'Proper train/test split, cross-validation, appropriate metrics for imbalanced data, avoiding leakage', 'Slide 6 (model comparison), Step 11 (CV tuning), leakage discovery story if asked'],
        ['Interpretability', 'Can you explain why the model makes its predictions?', 'Slide 9 (SHAP), feature engineering rationale, confusion matrix walkthrough'],
        ['Fairness/equity', 'Does the model work for everyone? Did you even check?', 'Slide 10 (subgroup analysis across 5 dimensions). This differentiates us from most teams.'],
        ['Presentation quality', 'Clear visuals, confident delivery, time management', 'Clean pres_ figures, practiced narrative arc, 1 min per slide'],
        ['Novelty/insight', 'Did you find something surprising or do something creative?', 'The confounding story, treatment_resistant feature, leakage discovery, age paradox, NPV as the star metric'],
        ['Practical deployment', 'Could this actually be used?', 'Slide 11 (threshold table, workflow), limitations slide'],
    ]
)

doc.add_heading('6.1 Slide Design Tips', level=2)

for tip in [
    'One message per slide. If a slide makes two points, split it into two slides.',
    'Figures should fill 60-70% of the slide. Titles and a few bullets, no paragraphs.',
    'Use the pres_ figures — they are already sized and styled for slides. Do not resize the step_ plots into tiny thumbnails.',
    'For tables (Slides 3, 5, 11), use your software\'s built-in table formatting. Keep fonts ≥ 18pt.',
    'Practice the transitions: "We saw that A1c dominates → but it is not enough → so we engineered features → here is how the model uses them."',
    'Have backup slides ready with step_ exploration plots in case judges ask for more detail.',
    'Time yourself: 10–12 minutes of content, leaving 3–5 minutes for Q&A.',
    'End on Slide 13 (Summary) with 5 crisp bullet points. Do not end on limitations.',
]:
    doc.add_paragraph(tip, style='List Bullet')

doc.add_heading('6.2 Narrative Arc', level=2)

doc.add_paragraph(
    'The presentation tells a story in three acts:'
)

bold_para('Act 1 — The Challenge (Slides 1–4, ~3 minutes):')
doc.add_paragraph(
    '"10.6% of diabetes patients lose glycemic control every year. A1c is the strongest predictor '
    'but misses 65% of cases. We need a multivariate model."'
)

bold_para('Act 2 — Our Approach (Slides 5–9, ~5 minutes):')
doc.add_paragraph(
    '"We engineered 20 clinically-driven features, tested 3 model families, tuned the best one. '
    'Tuned XGBoost: AUC 0.852, catches 81%, NPV 97%. SHAP shows it reasons like a clinician."'
)

bold_para('Act 3 — Impact & Trust (Slides 10–13, ~3 minutes):')
doc.add_paragraph(
    '"The model is equitable across all subgroups. It can be deployed as a monthly screening tool '
    'with adjustable thresholds. This enables proactive care — finding at-risk patients before '
    'they lose control."'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('Quick Reference Card', level=1)

bold_para('The 5 numbers you must know:')
add_table(
    ['Metric', 'Value', 'What to say'],
    [
        ['Patients', '62,425', '"Over sixty thousand diabetes patients"'],
        ['AUC-ROC', '0.852', '"Eighty-five percent discrimination"'],
        ['Sensitivity', '80.6%', '"Catches four out of five at-risk patients"'],
        ['NPV', '97.0%', '"When the model says low risk, it is right 97% of the time"'],
        ['Features', '48 (20 engineered)', '"Forty-eight features, twenty engineered from clinical reasoning"'],
    ]
)

doc.add_paragraph()
bold_para('Plot file quick-reference:')
add_table(
    ['Need', 'Use this file'],
    [
        ['Class balance for Slide 2', 'pres_1_problem.png'],
        ['A1c risk gradient for Slide 4', 'pres_2_risk_by_a1c.png'],
        ['ROC comparison for Slide 6', 'pres_3_model_roc.png'],
        ['Key metrics for Slide 7', 'pres_7_metrics.png'],
        ['Confusion matrix for Slide 8', 'pres_4_confusion.png'],
        ['SHAP for Slide 9', 'pres_5_shap.png'],
        ['Equity for Slide 10', 'pres_6_fairness.png'],
        ['Detailed evaluation (backup)', 'step12_full_evaluation.png'],
        ['Probability distribution (backup)', 'step12_probability_distribution.png'],
        ['Missingness (backup)', 'step4_missingness_by_outcome.png'],
        ['Full correlation matrix (backup)', 'step5_correlation_matrix.png'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save('Role_C_Data_Explorer_Presentation.docx')
print('Saved: Role_C_Data_Explorer_Presentation.docx')
