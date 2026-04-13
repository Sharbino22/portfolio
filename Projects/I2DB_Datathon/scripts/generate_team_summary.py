"""
Generate Team_Summary.docx -- 2-page executive summary for team meeting.
Covers: what we did, key decisions, results, what's new with test set,
role assignments, and what each person needs to know.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles['Heading {}'.format(level)]
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)
    h.font.size = Pt(14 - level)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

def bold_para(text, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('I2DB Datathon 2025 -- Team Briefing')
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Read this before the team meeting. Full details in datathon_notebook_final.docx.')
run.italic = True
run.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: THE BIG PICTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. What We Did', level=1)

doc.add_paragraph(
    'We built a machine learning model that predicts whether a diabetes patient\'s A1c will become '
    'uncontrolled in the next year, using 12 months of prior EHR data. We followed a 16-step pipeline:'
)

add_table(
    ['Phase', 'Steps', 'What happened'],
    [
        ['Explore', '1-5', 'Loaded 62,425 patients x 41 features. Found 10.6% uncontrolled rate (1:8 imbalance). A1c is the strongest predictor. Discovered medication confounding and a leakage column.'],
        ['Prepare', '6-8', 'Cleaned data (41 -> 36 columns). Engineered 20 new features based on clinical reasoning. Split 80/20 for internal validation.'],
        ['Model', '9-12', 'Tested logistic regression, random forest, XGBoost. Tuned XGBoost (AUC improved from 0.827 to 0.852). Evaluated on held-out 20%.'],
        ['Interpret', '13-14', 'SHAP analysis shows a1c_mean is the #1 predictor. Subgroup analysis shows no bias against minorities -- model actually performs BETTER for Black and Medicaid patients.'],
        ['Submit', '15-16', 'Received competition test set (15,607 patients). Retrained model on 100% of training data. Generated final predictions. Built presentation outline.'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: KEY RESULTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Key Results', level=1)

add_table(
    ['Metric', 'Value', 'What it means'],
    [
        ['AUC-ROC', '0.854', 'The model correctly ranks an uncontrolled patient as higher risk than a controlled patient 85% of the time.'],
        ['Sensitivity', '~78%', 'Catches roughly 4 out of 5 patients who will become uncontrolled.'],
        ['NPV', '97%', 'If the model says a patient is low-risk, 97% of the time it is correct.'],
        ['Equity', 'No disparities', 'AUC ranges 0.83-0.90 across all race, gender, age, insurance groups. Higher sensitivity for Black and Medicaid patients.'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: KEY DECISIONS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Key Decisions You Need to Understand', level=1)

doc.add_paragraph(
    'These are the decisions judges will ask about. Every team member should be able to explain these.'
)

bold_para('1. Why we dropped the leakage column')
doc.add_paragraph(
    'The column "a1c 2025 collection date" was 0% missing for uncontrolled patients but 47.7% '
    'missing for controlled ones. If we kept it, the model could cheat by checking if the column '
    'is missing. We dropped it because it would not exist in a real clinical deployment.'
)

bold_para('2. Why medications appear to INCREASE risk (confounding)')
doc.add_paragraph(
    'Patients on more medications have higher uncontrolled rates. This is NOT because medications '
    'cause poor control -- it is because sicker patients get prescribed more medications. This is '
    'reverse causation. We handled it by creating the "treatment_resistant" feature (A1c >= 8 despite '
    '2+ medication classes) which captures the real clinical signal.'
)

bold_para('3. Why we chose XGBoost over random forest')
doc.add_paragraph(
    'Random forest won with default settings (AUC 0.847), but XGBoost surpassed it after tuning '
    '(AUC 0.852). XGBoost learns sequentially -- each tree corrects the previous one\'s mistakes. '
    'Think of it like a tumor board where each specialist refines the diagnosis.'
)

bold_para('4. Why we retrained on 100% of the data')
doc.add_paragraph(
    'We used 80/20 split during development to validate our model honestly. Once we confirmed it '
    'works (AUC 0.854), we retrained on all 62,425 patients before predicting on the competition '
    'test set. This is standard practice and was explicitly recommended by Adam Wilcox.'
)

bold_para('5. Why we chose threshold 0.50')
doc.add_paragraph(
    'The model outputs a probability for each patient. We chose 0.50 as the cutoff because this is '
    'a screening tool -- missing a patient who will become uncontrolled (and develop complications '
    'like retinopathy, nephropathy) is far worse than scheduling an unnecessary follow-up. '
    'At 0.50, we catch ~78% of at-risk patients.'
)

bold_para('6. Feature engineering -- the most impactful step')
doc.add_paragraph(
    'We created 20 features from clinical reasoning. The top ones by importance (SHAP):'
)
add_table(
    ['Feature', 'What it captures', 'Why it matters'],
    [
        ['a1c_mean', 'Average of all A1c readings', '#1 predictor. Glycemic burden over time.'],
        ['a1c_max', 'Worst A1c reading ever', 'History of severe episodes signals ongoing risk.'],
        ['a1c_change', 'Latest A1c minus oldest A1c', 'Trajectory -- getting better or worse?'],
        ['treatment_resistant', 'A1c >= 8 despite 2+ med classes', 'Captures disease that doesn\'t respond to treatment.'],
        ['n_a1c_tests', 'How many A1c tests were done', 'Proxy for monitoring intensity / disease severity.'],
        ['undertreated', 'A1c >= 8 with zero medications', 'Patient who should be on meds but isn\'t.'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: WHAT'S NEW -- COMPETITION TEST SET
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. What Happened With the Competition Test Set', level=1)

doc.add_paragraph(
    'The competition organizers released 15,607 new patients (TEST SET DM Features.csv). We do NOT '
    'have the answers (True/False) for these patients -- Adam\'s team holds those. We submit our '
    'predictions and they grade it.'
)

add_table(
    ['What we checked', 'Result'],
    [
        ['Same columns?', 'Yes -- all 41 columns identical'],
        ['Patient overlap?', 'Zero -- completely separate from training'],
        ['Missingness patterns?', 'Nearly identical (all within 1 percentage point)'],
        ['Feature distributions?', 'Overlapping -- same patient population'],
        ['Subgroup patterns?', 'Preserved -- males > females, younger > older, same as training'],
    ]
)

bold_para('Submission: final_submission.csv')
doc.add_paragraph(
    '15,607 patients. True/False predictions. Format matches DM_Control_2025.csv. '
    'Must be submitted by April 6. If we make top 3, we present at the I2DB Symposium on April 13.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: ROLE ASSIGNMENTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Role Assignments for the Presentation', level=1)

doc.add_paragraph(
    'If we make top 3, each person presents specific slides. Read your role document for full details. '
    'Here is what each role covers and what you MUST be able to explain:'
)

add_table(
    ['Role', 'Covers', 'Must be able to explain', 'Key document'],
    [
        ['A: Clinical Lead',
         'Feature engineering rationale, clinical narrative, why features matter',
         'Medication confounding (reverse causation), A1c risk gradient, '
         'why 65% of future-uncontrolled have A1c <9, treatment_resistant logic',
         'Role_A_Clinical_Lead.docx'],
        ['B: Clinical Support',
         'Literature grounding, demographics, subgroup/equity analysis',
         'Why model performs BETTER for minorities, diabetes disparities literature, '
         'age paradox (younger = higher risk), Medicaid recall = 91-93%',
         'Role_B_Clinical_Support.docx'],
        ['C: Data Explorer',
         'Build slides, manage plot inventory, narrate the data story',
         'Every exploration finding (Steps 1-5), all 25 figures, 13-slide arc, '
         'class imbalance (1:8), missingness patterns, leakage detection',
         'Role_C_Data_Explorer_Presentation.docx'],
        ['D: Technical Lead',
         'Code pipeline, model decisions, SHAP, technical Q&A',
         'Why XGBoost > RF after tuning, hyperparameter rationale, threshold tradeoffs, '
         'SHAP vs gain importance, no test set contamination, AUC definition',
         'Role_D_Technical_Lead.docx'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6: WHAT TO READ
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. What to Read Before the Meeting', level=1)

add_table(
    ['Priority', 'Document', 'Time', 'Who should read it'],
    [
        ['Required', 'This summary (Team_Summary.docx)', '5-10 min', 'Everyone'],
        ['Required', 'Your role document (Role_A/B/C/D.docx)', '15-20 min', 'Your assigned role'],
        ['Recommended', 'Full notebook (datathon_notebook_final.docx)', '45-60 min', 'Everyone (at least skim Sections 1, 6, 7)'],
        ['Reference', 'Presentation outline (PRESENTATION_OUTLINE.md)', '10 min', 'Role C especially, but everyone should see the slide structure'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7: TIMELINE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Timeline', level=1)

add_table(
    ['Date', 'What'],
    [
        ['Now', 'Team reads this summary + their role document'],
        ['Team meeting', 'Walk through the pipeline, assign final roles, discuss any questions'],
        ['Before April 6', 'Submit final_submission.csv to competition portal'],
        ['After April 6', 'Competition grades submissions, announces top 3'],
        ['April 13', 'I2DB Symposium -- presentation (if top 3)'],
    ]
)

doc.add_paragraph()
bold_para('Questions? Concerns? Bring them to the meeting.')

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save('docs/Team_Summary.docx')
print('Saved: docs/Team_Summary.docx')
