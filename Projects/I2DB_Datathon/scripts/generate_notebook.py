"""
Generate datathon_notebook.docx — comprehensive technical reference document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Style setup ───────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)

def add_image(path, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f'[Image not found: {path}]')

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('I2DB Datathon 2025')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Diabetes Risk Prediction')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Technical Notebook')
run.font.size = Pt(16)
run.italic = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Team: [Your Name(s) Here]')
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Institution: [Your Institution]')
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('I2DB Symposium — April 13, 2026')
run.font.size = Pt(14)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. Data Exploration (Steps 1–5)',
    '   2.1 Step 1: Loading and Initial Inspection',
    '   2.2 Step 2: Target Variable Analysis',
    '   2.3 Step 3: Feature Group Exploration',
    '   2.4 Step 4: Missingness Analysis',
    '   2.5 Step 5: Correlations',
    '3. Data Preparation (Steps 6–8)',
    '   3.1 Step 6: Data Cleaning',
    '   3.2 Step 7: Feature Engineering',
    '   3.3 Step 8: Train-Test Split',
    '4. Modeling (Steps 9–12)',
    '   4.1 Step 9: Baseline Logistic Regression',
    '   4.2 Step 10: Advanced Models',
    '   4.3 Step 11: Hyperparameter Tuning',
    '   4.4 Step 12: Test Set Evaluation',
    '5. Interpretation (Steps 13–14)',
    '   5.1 Step 13: Feature Importance and SHAP',
    '   5.2 Step 14: Subgroup Fairness Analysis',
    '6. Submission and Presentation',
    '7. Key Decisions Log',
    '8. Glossary',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith('   '):
        for run in p.runs:
            run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)

doc.add_heading('The Clinical Problem', level=2)
doc.add_paragraph(
    'Diabetes affects over 37 million Americans, and glycemic control — measured by hemoglobin A1c — '
    'is the cornerstone of diabetes management. When a patient\'s A1c becomes "uncontrolled" (generally '
    'A1c ≥ 9%), the risk of complications rises sharply: kidney disease, retinopathy, neuropathy, and '
    'cardiovascular events. Currently, clinicians typically discover loss of glycemic control reactively — '
    'the patient comes in for a routine visit, labs are drawn, and the A1c comes back elevated. By then, '
    'the patient may have been uncontrolled for months.'
)
doc.add_paragraph(
    'What if we could identify patients who are about to lose control before it happens? A predictive '
    'model could flag high-risk patients for proactive intervention: closer monitoring, medication '
    'adjustment, care coordination, or outreach. This is the clinical problem our project addresses.'
)

doc.add_heading('Competition Overview', level=2)
add_table(
    ['Item', 'Detail'],
    [
        ['Competition', 'I2DB Datathon 2025'],
        ['Prize', '$2,000 (first place)'],
        ['Submission deadline', 'April 6, 2026'],
        ['Symposium presentation', 'April 13, 2026'],
        ['Task', 'Binary classification — predict a1c 2025 Uncontrolled (True/False)'],
        ['Evaluation', 'Model performance, clinical reasoning, presentation quality'],
    ]
)

doc.add_heading('Our Approach in Plain English', level=2)
doc.add_paragraph(
    'We used 12 months of electronic health record (EHR) data for 62,425 diabetes patients to predict '
    'which patients will have uncontrolled A1c in the following year. Our approach followed a systematic '
    '16-step pipeline: explore the data, clean it, engineer clinically meaningful features, train '
    'multiple machine learning models, tune the best one, and evaluate it fairly across demographic '
    'subgroups. The final model — a tuned XGBoost classifier — achieves an AUC of 0.852 and catches '
    '81% of patients who will become uncontrolled.'
)

doc.add_heading('Dataset Overview', level=2)
add_table(
    ['Property', 'Value'],
    [
        ['Patients', '62,425'],
        ['Feature columns', '41 (demographics, labs, medications, utilization, comorbidities, social determinants)'],
        ['Target variable', 'a1c 2025 Uncontrolled (True = 10.6%, False = 89.4%)'],
        ['Feature file', 'DM_Features.csv'],
        ['Target file', 'DM_Control_2025.csv'],
        ['Class imbalance', '~1:8 ratio (uncontrolled : controlled)'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: DATA EXPLORATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Data Exploration (Steps 1–5)', level=1)

# --- Step 1 ---
doc.add_heading('2.1 Step 1: Loading and Initial Inspection', level=2)

bold_para('What we did:')
doc.add_paragraph(
    'We loaded both CSV files, verified they contained the same 62,425 patients in the same order, '
    'examined column names, data types, and missing values across all 41 features.'
)

bold_para('What we found:')
doc.add_paragraph('The 41 features fell into 9 clinical groups:')
add_table(
    ['Group', 'Columns', 'Key observations'],
    [
        ['Demographics', '4', 'Birth year (integer), sex, race, ethnicity. Near-complete.'],
        ['Comorbidities', '2', 'CAD and COPD counts. Zero missingness.'],
        ['A1c Labs', '11', 'Up to 5 A1c results + timing columns + 1 leakage column. A1c 1 is 100% complete; A1c 4-5 are 97% missing.'],
        ['Weight / Height', '6', 'Values, dates, units. ~73-76% missing.'],
        ['Cholesterol', '6', 'LDL, HDL, total + timing. ~24% missing.'],
        ['Utilization', '3', 'ED visits, PCP visits, admissions. Zero missingness.'],
        ['Insurance', '1', '67% missing.'],
        ['Medications', '6', 'Order counts for 6 drug classes. Zero missingness (0 = no orders).'],
        ['ADI', '2', 'Area Deprivation Index. 37-48% missing. Stored as strings — needed type conversion.'],
    ]
)

bold_para('What this means:')
doc.add_paragraph(
    'The data has a clear structure: some features are fully populated (medications, utilization, '
    'comorbidities, most demographics) while others are sparse (later A1c values, weight/height, '
    'insurance). The missingness pattern itself carries information — we return to this in Step 4. '
    'We also flagged the column "a1c 2025-collection date-time-days from reference" as a potential '
    'data leakage risk, which we investigate thoroughly in Step 4.'
)

# --- Step 2 ---
doc.add_heading('2.2 Step 2: Target Variable Analysis', level=2)

bold_para('What we did:')
doc.add_paragraph(
    'We examined the class balance of the target variable, visualized the distribution of the most '
    'recent A1c value (A1c 1) split by outcome, and computed uncontrolled rates across A1c bins.'
)

bold_para('What we found:')
doc.add_paragraph(
    'The target is heavily imbalanced: 55,793 patients (89.4%) remain controlled, while only 6,632 '
    '(10.6%) become uncontrolled. This is roughly a 1:8 ratio. A model that simply predicts '
    '"controlled" for every patient would be 89.4% accurate but clinically useless — it would miss '
    'every at-risk patient.'
)

add_image('plots/step2_class_balance.png', 5.5)
doc.add_paragraph('Figure 1: Class distribution showing the 1:8 imbalance between controlled and uncontrolled patients.')

doc.add_paragraph(
    'The most recent A1c value differs substantially between groups. Controlled patients have a '
    'mean A1c of 6.92 (median 6.5), while uncontrolled patients have a mean of 8.78 (median 8.3) — '
    'a difference of nearly 2 full A1c points.'
)

add_image('plots/step2_a1c1_by_outcome.png', 5.5)
doc.add_paragraph('Figure 2: A1c distribution (histogram and box plot) by future outcome. The uncontrolled group has a higher and more spread distribution.')

doc.add_paragraph(
    'When we binned patients by current A1c level, a clear dose-response emerged:'
)

add_table(
    ['Current A1c', 'Risk of becoming uncontrolled', 'N patients'],
    [
        ['< 5.7', '1.0%', '10,469'],
        ['5.7 – 6.5', '2.3%', '18,684'],
        ['6.5 – 7.0', '6.5%', '9,676'],
        ['7.0 – 8.0', '15.7%', '11,513'],
        ['8.0 – 9.0', '27.8%', '5,078'],
        ['9.0 – 10.0', '32.1%', '2,591'],
        ['> 10.0', '32.0%', '4,414'],
    ]
)

add_image('plots/step2_uncontrolled_rate_by_a1c_bin.png', 5.5)
doc.add_paragraph('Figure 3: Risk of becoming uncontrolled by current A1c level. Risk plateaus above A1c 9.')

bold_para('What this means:')
doc.add_paragraph(
    'Current A1c is the single strongest predictor, but it is not sufficient. Risk plateaus above '
    'A1c 9, suggesting that once a patient is severely uncontrolled, factors other than the A1c '
    'number itself (medications, engagement, comorbidities) determine whether they stay uncontrolled. '
    'Critically, 65% of patients who become uncontrolled currently have A1c below 9. The model needs '
    'features beyond A1c to catch these patients.'
)

# --- Step 3 ---
doc.add_heading('2.3 Step 3: Feature Group Exploration', level=2)

bold_para('What we did:')
doc.add_paragraph(
    'We explored each feature group — demographics, comorbidities, medications, utilization, '
    'cholesterol, and ADI — by comparing distributions and uncontrolled rates between groups.'
)

bold_para('Demographics:')
doc.add_paragraph(
    'Younger patients have higher uncontrolled rates: 14.3% for patients under 30 versus 8.8% for '
    'those over 80. Males (11.8%) are slightly more at risk than females (9.6%). Race shows '
    'meaningful disparities: Pacific Islander (15.4%) and American Indian (14.9%) have the highest '
    'rates, while White (10.1%) and Asian (10.2%) have the lowest. Hispanic patients (12.8%) are '
    'higher than non-Hispanic (10.6%).'
)

add_image('plots/step3_demographics.png', 5.5)
doc.add_paragraph('Figure 4: Uncontrolled rates by age group, gender, and race.')

bold_para('Medications (important — confounding):')
doc.add_paragraph(
    'Uncontrolled patients use more medications, not fewer. Insulin use is 44.7% in uncontrolled '
    'versus 32.8% in controlled (+12 percentage points). Metformin and sulfonylurea show similar '
    'gaps (+10-11pp). This seems counterintuitive — shouldn\'t more treatment lead to better control?'
)
doc.add_paragraph(
    'This is confounding due to reverse causation. Patients with high A1c get prescribed more '
    'medications, not the other way around. A patient on 3 drugs with A1c 7.5 is a fundamentally '
    'different patient from one on 0 drugs with A1c 7.5 — the first has aggressive, treatment-'
    'resistant disease that required 3 classes to get down to 7.5, while the second has mild disease '
    'managed by diet alone. This insight directly informed our feature engineering: we created '
    '"treatment_resistant" (high A1c despite multiple medications) to properly capture this signal.'
)

add_image('plots/step3_comorbidities_meds.png', 5.5)
doc.add_paragraph('Figure 5: Comorbidities and medication use by outcome. Note higher medication use in the uncontrolled group.')

bold_para('Utilization, Cholesterol, and ADI:')
doc.add_paragraph(
    'ED visits showed a dose-response: 9.7% uncontrolled at 0 visits rising to 18.4% at 5+ visits. '
    'PCP visits and admissions showed minimal differences. Cholesterol (LDL, HDL, total) and ADI '
    'national rank were nearly identical between groups — weak individual predictors, though they '
    'may contribute in combination with other features inside a tree-based model.'
)

add_image('plots/step3_util_chol_adi.png', 5.5)
doc.add_paragraph('Figure 6: Utilization, cholesterol, and ADI distributions by outcome.')

bold_para('Feature signal ranking:')
doc.add_paragraph(
    'Based on exploration, we ranked feature groups by predictive signal strength: '
    'A1c values (strongest) >> medications (strong, confounded) > demographics (moderate) > '
    'ED visits (moderate) >> cholesterol and ADI (weak individually).'
)

# --- Step 4 ---
doc.add_heading('2.4 Step 4: Missingness Analysis', level=2)

bold_para('What we did:')
doc.add_paragraph(
    'We mapped which columns have missing data, visualized the missingness pattern, and — crucially — '
    'tested whether missingness rates differ between controlled and uncontrolled patients.'
)

bold_para('What we found:')
doc.add_paragraph(
    'Missingness is not random. It differs by outcome for several important columns:'
)

add_table(
    ['Column', 'Controlled % missing', 'Uncontrolled % missing', 'Difference'],
    [
        ['a1c 2025 (date) — LEAKAGE', '47.7%', '0.0%', '-47.7pp'],
        ['A1c 2', '59.7%', '40.9%', '-18.8pp'],
        ['A1c 3', '89.1%', '76.3%', '-12.8pp'],
        ['Cholesterol (HDL/total)', '25.1%', '20.0%', '-5.1pp'],
        ['Weight/height', '~73-76%', '~74-76%', '~0pp (similar)'],
        ['Insurance', '67.1%', '66.2%', '~0pp (similar)'],
    ]
)

add_image('plots/step4_missingness_by_outcome.png', 5.5)
doc.add_paragraph('Figure 7: Missingness rates compared between controlled and uncontrolled patients.')

add_image('plots/step4_missingness_heatmap.png', 5.5)
doc.add_paragraph('Figure 8: Missingness heatmap for 500 random patients. Red = missing, green = present.')

bold_para('The leakage column:')
doc.add_paragraph(
    'The column "a1c 2025-collection date-time-days from reference" is 0% missing for uncontrolled '
    'patients but 47.7% missing for controlled ones. This means if this column is missing, the '
    'patient is always controlled. A model could "cheat" by simply checking whether this column is '
    'null. We confirmed this with a cross-tabulation: 100% of the 26,617 patients missing this '
    'value are in the controlled group. This column was dropped immediately in data cleaning.'
)

bold_para('What this means:')
doc.add_paragraph(
    'Missingness itself carries information. Uncontrolled patients had more A1c tests (A1c 2 and 3 '
    'are much less missing), which makes clinical sense — if your A1c is high, your doctor orders '
    'repeat labs more frequently. This insight informed our "n_a1c_tests" feature: the count of '
    'non-missing A1c values captures monitoring intensity, which is a proxy for disease severity.'
)

# --- Step 5 ---
doc.add_heading('2.5 Step 5: Correlations', level=2)

bold_para('What we did:')
doc.add_paragraph(
    'We computed Pearson correlations between all numeric features and the target variable, '
    'and identified highly correlated feature pairs that might be redundant.'
)

bold_para('Correlations with the target:')
add_table(
    ['Feature', 'Correlation', 'Strength'],
    [
        ['A1c 5 (oldest reading)', '+0.432', 'Strong'],
        ['A1c 4', '+0.404', 'Strong'],
        ['A1c 3', '+0.387', 'Strong'],
        ['A1c 2', '+0.386', 'Strong'],
        ['A1c 1 (most recent)', '+0.314', 'Strong'],
        ['Sulfonylurea orders', '+0.109', 'Moderate'],
        ['Insulin orders', '+0.082', 'Moderate'],
        ['Metformin orders', '+0.066', 'Moderate'],
        ['All others', '< 0.05', 'Weak'],
    ]
)

add_image('plots/step5_target_correlations.png', 5.5)
doc.add_paragraph('Figure 9: Feature correlations with the target, ranked by magnitude.')

bold_para('Redundant feature pairs (|r| > 0.5):')
add_table(
    ['Feature A', 'Feature B', 'Correlation', 'Action taken'],
    [
        ['ADI state rank', 'ADI national rank', '+0.964', 'Dropped state (kept national)'],
        ['LDL', 'Total cholesterol', '+0.772', 'Kept both (model can sort it out)'],
        ['PCP visits', 'Admissions', '+0.733', 'Engineered total_encounters'],
        ['A1c 1', 'A1c 2', '+0.642', 'Aggregated into summary features'],
        ['Various A1c pairs', 'Various', '0.52–0.56', 'Aggregated into mean, max, change, variability'],
    ]
)

add_image('plots/step5_correlation_matrix.png', 5.5)
doc.add_paragraph('Figure 10: Full correlation matrix. Note the A1c cluster (upper left) and cholesterol cluster.')

bold_para('What this means:')
doc.add_paragraph(
    'A1c values dominate predictive signal. The five raw A1c columns correlate +0.31 to +0.43 with '
    'the target — far stronger than any other feature. Everything else (medications, demographics, '
    'utilization) adds marginal signal individually. This guided our decision to aggregate the raw '
    'A1c columns into summary features (mean, max, change, variability) rather than using them directly, '
    'and to drop the redundant ADI state rank.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: DATA PREPARATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Data Preparation (Steps 6–8)', level=1)

# --- Step 6 ---
doc.add_heading('3.1 Step 6: Data Cleaning', level=2)

doc.add_paragraph(
    'We transformed the raw 41-column dataset into a clean 36-column numeric dataset ready for '
    'feature engineering. Every column is now numeric (int64 or float64).'
)

bold_para('What we dropped (14 columns removed):')
add_table(
    ['Action', 'Columns', 'Reason'],
    [
        ['Dropped leakage column', '1', '100% of missing values = Controlled. Would allow the model to cheat.'],
        ['Dropped date/timing columns', '10', 'When a lab was collected is less useful than the result itself. Timing is captured indirectly by n_a1c_tests.'],
        ['Dropped unit columns', '2', 'All weights are kg, all heights are cm. Uniform units = zero information.'],
        ['Dropped ADI state rank', '1', '96% correlated with ADI national rank. Keeping both would be redundant.'],
    ]
)

bold_para('What we encoded (4 text columns → 13 numeric columns):')
add_table(
    ['Original column', 'Encoding', 'Result'],
    [
        ['Gender at birth', 'Binary: is_male (0/1)', '1 column'],
        ['Ethnicity', 'Binary: is_hispanic (0/1)', '1 column'],
        ['Race', 'One-hot encoding', '7 columns (one per category)'],
        ['Insurance', 'Simplified from 12 categories to 4, then one-hot', '4 columns (Managed Care, Medicaid, Medicare, Other)'],
    ]
)

doc.add_paragraph(
    'One-hot encoding means creating a separate binary (0/1) column for each category. For example, '
    'a patient who is Black and on Medicare would have race_Black_or_African_American = 1 (all other '
    'race columns = 0) and insurance_Medicare = 1 (all other insurance columns = 0). This allows '
    'the model to learn different effects for each category.'
)

bold_para('What we fixed:')
doc.add_paragraph(
    'ADI national rank was stored as a string (e.g., "10.0") instead of a number. We converted it '
    'to float64 using pd.to_numeric() with errors="coerce" to handle any non-numeric values.'
)

# --- Step 7 ---
doc.add_heading('3.2 Step 7: Feature Engineering', level=2)

doc.add_paragraph(
    'Feature engineering is the process of creating new, more informative variables from the raw '
    'data. This was the most impactful step in our pipeline — the engineered features capture '
    'clinical reasoning that raw columns cannot express. We created 20 new features organized into '
    '6 categories, guided by both data science best practices and clinical domain knowledge.'
)

bold_para('A1c-Derived Features (8 features):')
doc.add_paragraph(
    'These are the most important features in the model. Rather than feeding in up to 5 raw A1c '
    'readings (most of which are missing for most patients), we summarized them into meaningful '
    'clinical metrics.'
)

add_table(
    ['Feature', 'Calculation', 'Clinical reasoning'],
    [
        ['age', '2025 − birth year', 'Younger patients showed higher uncontrolled rates in Step 3 (14.3% for <30 vs 8.8% for 80+).'],
        ['a1c_latest', 'A1c 1 (most recent value)', 'The current A1c is the single strongest raw predictor (r = +0.31 with target).'],
        ['a1c_mean', 'Mean of all available A1c readings', 'Smooths out single-test noise. Captures overall glycemic burden. Strongest engineered feature (SHAP = 0.73).'],
        ['a1c_max', 'Maximum A1c across all readings', 'Was the patient ever severely uncontrolled? A past episode of A1c 12 is informative even if current A1c is 7.'],
        ['a1c_change', 'Most recent A1c − oldest A1c', 'Trajectory: positive means A1c went up (worsening), negative means improving. Only computed for patients with 2+ tests.'],
        ['a1c_variability', 'Standard deviation of A1c readings', 'Glycemic instability. A patient who swings between 6 and 10 is different from one stable at 8, even though both average 8.'],
        ['n_a1c_tests', 'Count of non-missing A1c values (1–5)', 'More monitoring = higher-risk patient. 58% had only 1 test; patients with 3+ tests were being watched closely.'],
        ['a1c_above_9', '1 if latest A1c ≥ 9, else 0', 'Clinical threshold for severe uncontrolled diabetes. The #1 feature by XGBoost built-in importance.'],
    ]
)

bold_para('Medication Features (4 features):')
add_table(
    ['Feature', 'Calculation', 'Clinical reasoning'],
    [
        ['total_med_classes', 'Count of medication types with orders > 0', 'Treatment intensity. Ranges from 0 (diet-controlled) to 6 (on everything).'],
        ['total_med_orders', 'Sum of all medication order counts', 'Total prescription volume, not just number of drug types.'],
        ['on_insulin', '1 if insulin orders > 0', '34% of patients. Insulin = more advanced disease requiring injectable therapy.'],
        ['on_newer_drugs', '1 if GLP-1 or SGLT2 orders > 0', '13% of patients. These newer, expensive drugs are often reserved for harder-to-treat cases.'],
    ]
)

bold_para('Treatment Pattern Features (3 features):')
doc.add_paragraph(
    'These features capture the relationship between disease severity and treatment — addressing '
    'the medication confounding we discovered in Step 3.'
)
add_table(
    ['Feature', 'Calculation', 'Clinical reasoning'],
    [
        ['no_medication', '1 if total_med_classes = 0', '42% of patients are on zero diabetes medications. These are typically mild, diet-controlled patients with low risk.'],
        ['undertreated', '1 if A1c ≥ 8 AND total_med_classes = 0', 'The patient who SHOULD be on medication but is not. 4.9% of patients. Could indicate non-adherence, access barriers, or clinical inertia.'],
        ['treatment_resistant', '1 if A1c ≥ 8 AND total_med_classes ≥ 2', 'The patient who IS on multiple meds but still cannot control their A1c. 6.4% of patients. One of the strongest non-A1c predictors (r = +0.20).'],
    ]
)

bold_para('Body Metrics (2 features):')
add_table(
    ['Feature', 'Calculation', 'Clinical reasoning'],
    [
        ['bmi', 'Weight (kg) / (height (m))²', 'Body mass index. Mean 32.6 (obese range). Only available for 23.5% of patients due to missingness in weight/height.'],
        ['has_bmi', '1 if BMI can be calculated', 'Whether weight and height were recorded. Missingness itself may be informative — patients seen more frequently may be more likely to have vitals documented.'],
    ]
)

bold_para('Utilization Features (3 features):')
add_table(
    ['Feature', 'Calculation', 'Clinical reasoning'],
    [
        ['total_encounters', 'ED visits + PCP visits + admissions', 'Total healthcare utilization. Combines three correlated columns (PCP and admissions were r = 0.73).'],
        ['any_admission', '1 if admission count > 0', '32% of patients had at least one hospitalization.'],
        ['any_comorbidity', '1 if CAD > 0 or COPD > 0', '23% have at least one comorbidity. Captures overall disease burden.'],
    ]
)

bold_para('Feature validation:')
doc.add_paragraph(
    'After engineering all 20 features, we computed their correlations with the target to confirm '
    'they carry signal. The top engineered features by correlation were: a1c_max (+0.35), a1c_mean '
    '(+0.33), a1c_latest (+0.31), a1c_above_9 (+0.25), a1c_variability (+0.21), and '
    'treatment_resistant (+0.20). The treatment_resistant feature was especially validating — it '
    'confirmed our hypothesis from Step 3 that combining A1c level with medication count captures '
    'a clinically meaningful signal that neither variable captures alone.'
)

# --- Step 8 ---
doc.add_heading('3.3 Step 8: Train-Test Split', level=2)

doc.add_paragraph(
    'We split the 62,425 patients into a training set (80%) and a held-out test set (20%) using '
    'scikit-learn\'s train_test_split with random_state=42 for reproducibility.'
)

add_table(
    ['Set', 'Patients', 'Controlled', 'Uncontrolled', '% Uncontrolled'],
    [
        ['Train', '49,940', '44,634', '5,306', '10.62%'],
        ['Test', '12,485', '11,159', '1,326', '10.62%'],
    ]
)

doc.add_paragraph(
    'We used stratified splitting, which ensures the class balance (10.62% uncontrolled) is '
    'preserved in both sets. Without stratification, random chance could put more uncontrolled '
    'patients in one set than the other, which would make our evaluation unreliable. It is the '
    'same principle as ensuring both arms of a clinical trial have the same disease prevalence.'
)

doc.add_paragraph(
    'We also dropped 8 raw columns that were now fully captured by engineered features: the 5 raw '
    'A1c value columns, birth year (replaced by age), and weight/height (replaced by bmi and '
    'has_bmi). The final feature matrix contains 48 columns.'
)

bold_para('The test set is a sealed envelope.')
doc.add_paragraph(
    'The test set was not used for any training, feature selection, or tuning decisions. It was only '
    'opened once, in Step 12, to get our final unbiased performance estimate. This prevents '
    '"overfitting to the test set" — where iteratively tweaking a model to improve test performance '
    'gives an overly optimistic estimate that would not generalize to new patients.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: MODELING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Modeling (Steps 9–12)', level=1)

# --- Step 9 ---
doc.add_heading('4.1 Step 9: Baseline Logistic Regression', level=2)

bold_para('Why start with logistic regression?')
doc.add_paragraph(
    'Logistic regression is the simplest classification model — it finds a straight line (in high-'
    'dimensional space) that separates the two classes. It is the "vital signs check" before ordering '
    'advanced imaging. If logistic regression does well, the signal is strong and linear. If it does '
    'poorly, we know we need more complex approaches.'
)

doc.add_paragraph(
    'We used class_weight="balanced" to handle the 1:8 class imbalance. This tells the model to '
    'pay 8× more attention to misclassifying an uncontrolled patient than a controlled one. Without '
    'this, the model would simply predict "controlled" for everyone (89.4% accuracy, 0% usefulness).'
)

bold_para('Results:')
add_table(
    ['Model', 'AUC-ROC', 'PR-AUC', 'Recall', 'Precision'],
    [
        ['HIGH features only (6)', '0.8267', '0.3120', '72%', '28%'],
        ['ALL features (48)', '0.8275', '0.3216', '75%', '28%'],
        ['Dummy (always Controlled)', '0.5000', '—', '0%', '—'],
    ]
)

doc.add_paragraph(
    'Key finding: adding 42 more features barely improved AUC (+0.001). This tells us that in a '
    'linear model, A1c features carry almost all the signal. The medication and demographic features '
    'don\'t help much when you can only add them linearly. Tree-based models, which can discover '
    'interactions (e.g., high A1c AND many meds), should do better.'
)

add_image('plots/step9_baseline_roc_pr.png', 5.5)
doc.add_paragraph('Figure 11: ROC and precision-recall curves for the baseline logistic regression models.')

# --- Step 10 ---
doc.add_heading('4.2 Step 10: Advanced Models', level=2)

doc.add_paragraph(
    'We trained two tree-based models: Random Forest and XGBoost. Tree-based models work by asking '
    'a sequence of yes/no questions about the features (e.g., "Is A1c > 8? → Yes → Is the patient '
    'on insulin? → No → High risk"). They naturally capture interactions between features without '
    'needing to explicitly engineer every combination.'
)

bold_para('Results (all using default settings):')
add_table(
    ['Model', 'AUC-ROC', 'PR-AUC', 'Recall', 'Precision'],
    [
        ['Logistic Regression', '0.8275', '0.3216', '75.0%', '27.7%'],
        ['Random Forest', '0.8467', '0.3923', '68.9%', '31.5%'],
        ['XGBoost (defaults)', '0.8267', '0.3596', '64.8%', '30.4%'],
    ]
)

doc.add_paragraph(
    'Random Forest won with default settings (AUC 0.847), while XGBoost with defaults (AUC 0.827) '
    'was no better than logistic regression. This is expected — XGBoost is a more powerful algorithm '
    'but requires tuning to reach its potential, like a race car that needs the right setup for each '
    'track. Step 11 addresses this.'
)

add_image('plots/step10_model_comparison.png', 5.5)
doc.add_paragraph('Figure 12: ROC and PR curves comparing all three models. Random Forest (orange) leads with default settings.')

# --- Step 11 ---
doc.add_heading('4.3 Step 11: Hyperparameter Tuning', level=2)

doc.add_paragraph(
    'Hyperparameters are settings that control how the model learns — they are not learned from the '
    'data but set by us before training. Tuning them is like adjusting the settings on a microscope '
    'to get the sharpest image.'
)

doc.add_paragraph(
    'We used RandomizedSearchCV with 80 parameter combinations × 5-fold cross-validation = 400 '
    'total model fits. Each combination was evaluated on its cross-validated AUC-ROC, ensuring we '
    'selected parameters that generalize rather than overfit.'
)

bold_para('Final tuned parameters:')
add_table(
    ['Parameter', 'Value', 'Plain English explanation'],
    [
        ['n_estimators', '800', 'Number of trees. More trees = more chances to learn patterns.'],
        ['learning_rate', '0.01', 'How much each tree adjusts the prediction. Small = careful, gradual learning (usually better but slower).'],
        ['max_depth', '5', 'Maximum depth of each tree (how many questions in sequence). Limits complexity to prevent overfitting.'],
        ['min_child_weight', '10', 'A leaf node must contain at least 10 patients. Prevents the model from making decisions based on tiny groups.'],
        ['subsample', '0.7', 'Each tree sees 70% of training patients (randomly sampled). Like training residents on rotating patient panels — prevents memorization.'],
        ['colsample_bytree', '0.5', 'Each tree sees 50% of features. Forces the model to find multiple pathways to the answer, not just rely on A1c every time.'],
        ['gamma', '0.3', 'Minimum improvement required to make a split. Like requiring a minimum clinical significance before acting on a result.'],
        ['reg_lambda', '1', 'L2 regularization penalty. Discourages extreme weights on any single feature.'],
        ['scale_pos_weight', '8.41', 'Ratio of controlled to uncontrolled patients (44,634 / 5,306). Tells the model that missing an uncontrolled patient is 8.4× worse than a false alarm.'],
    ]
)

bold_para('Improvement from tuning:')
add_table(
    ['Metric', 'XGBoost (default)', 'XGBoost (tuned)', 'Improvement'],
    [
        ['AUC-ROC', '0.8267', '0.8517', '+0.025'],
        ['PR-AUC', '0.3596', '0.4162', '+0.057'],
        ['Recall', '64.8%', '80.6%', '+15.8pp'],
    ]
)

add_image('plots/step11_tuned_comparison.png', 5.5)
doc.add_paragraph('Figure 13: The tuned XGBoost (thick blue) clearly separates from all other models in both ROC and PR curves.')

# --- Step 12 ---
doc.add_heading('4.4 Step 12: Test Set Evaluation', level=2)

doc.add_paragraph(
    'This is the moment of truth — evaluating our final tuned model on the held-out test set of '
    '12,485 patients that the model has never seen during training or tuning.'
)

bold_para('Core metrics:')
add_table(
    ['Metric', 'Value', 'Clinical interpretation'],
    [
        ['AUC-ROC', '0.852', 'If you pick a random controlled and a random uncontrolled patient, the model correctly ranks the uncontrolled patient as higher risk 85% of the time.'],
        ['Sensitivity (Recall)', '80.6%', 'Of 1,326 patients who WILL become uncontrolled, the model catches 1,069 (4 out of 5).'],
        ['Specificity', '74.2%', 'Of 11,159 controlled patients, the model correctly clears 8,275.'],
        ['PPV (Precision)', '27.0%', 'Of every 4 patients flagged, 1 is truly uncontrolled. Expected for screening with a 10.6% base rate.'],
        ['NPV', '97.0%', 'Of patients the model clears as low-risk, 97% truly stay controlled. The strongest clinical metric.'],
        ['Brier Score', '0.159', 'Measures calibration (accuracy of probability estimates). Lower is better; 0 = perfect.'],
    ]
)

add_image('plots/pres_4_confusion.png', 5.5)
doc.add_paragraph('Figure 14: Confusion matrix with plain-language labels for each quadrant.')

add_image('plots/step12_full_evaluation.png', 5.5)
doc.add_paragraph('Figure 15: Four-panel evaluation: confusion matrix, ROC curve, PR curve, and calibration curve.')

bold_para('Threshold analysis:')
doc.add_paragraph(
    'The model outputs a probability (0–100%) for each patient. We choose a threshold to convert '
    'this to a yes/no decision. Different thresholds suit different clinical contexts:'
)

add_table(
    ['Threshold', 'Recall', 'Precision', 'Flagged (per 12,485)', 'Missed'],
    [
        ['0.15 (aggressive screening)', '96.4%', '16.6%', '7,697', '48'],
        ['0.30 (panel management)', '90.6%', '21.6%', '5,552', '125'],
        ['0.50 (default)', '80.6%', '27.0%', '3,953', '257'],
    ]
)

add_image('plots/step12_probability_distribution.png', 5.5)
doc.add_paragraph(
    'Figure 16: Distribution of predicted probabilities by true outcome. Green (controlled) '
    'clusters left, red (uncontrolled) clusters right. Good separation with some overlap in '
    'the 0.2–0.5 range.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: INTERPRETATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Interpretation (Steps 13–14)', level=1)

# --- Step 13 ---
doc.add_heading('5.1 Step 13: Feature Importance and SHAP', level=2)

doc.add_paragraph(
    'Understanding why a model makes its predictions is as important as the predictions themselves. '
    'A "black box" model that flags patients without explanation would not be trusted by clinicians. '
    'We used two complementary methods to interpret the model.'
)

bold_para('Method 1: XGBoost built-in importance (gain)')
doc.add_paragraph(
    'This measures how much each feature improves predictions when used in a tree split. It answers: '
    '"Which feature does the model check most often and benefit from most?"'
)

add_image('plots/step13_feature_importance.png', 5.0)
doc.add_paragraph('Figure 17: Top 20 features by XGBoost gain. a1c_above_9 dominates as a single high-information split.')

bold_para('Method 2: SHAP values')
doc.add_paragraph(
    'SHAP (SHapley Additive exPlanations) is a more rigorous method borrowed from game theory. '
    'For each patient and each feature, SHAP computes exactly how much that feature pushed the '
    'prediction toward "uncontrolled" or toward "controlled." The SHAP beeswarm plot below shows '
    'this for all test patients simultaneously.'
)

add_image('plots/step13_shap_summary.png', 5.5)
doc.add_paragraph(
    'Figure 18: SHAP beeswarm plot. Each dot is one patient. X-axis = impact on prediction. '
    'Color = feature value (red = high, blue = low).'
)

bold_para('How to read the SHAP plot:')
doc.add_paragraph(
    'Look at the top row (a1c_mean). Red dots (high A1c mean) are pushed to the right (toward '
    'uncontrolled). Blue dots (low A1c mean) are pushed to the left (toward controlled). This '
    'confirms: high average A1c is the single strongest driver of risk.'
)

bold_para('Top features and their clinical interpretation:')
add_table(
    ['Rank', 'Feature', 'Mean |SHAP|', 'Clinical interpretation'],
    [
        ['1', 'a1c_mean', '0.728', 'Average A1c is the dominant predictor. Higher mean = higher risk. This is the glycemic burden over the observation period.'],
        ['2', 'a1c_max', '0.501', 'The worst A1c reading matters — even if current A1c improved, a history of severe episodes increases risk.'],
        ['3', 'a1c_latest', '0.218', 'Most recent reading. Adds recency information beyond the mean.'],
        ['4', 'a1c_change', '0.108', 'Trajectory. Positive change (A1c went up) pushes toward uncontrolled. Negative (improving) is protective.'],
        ['5', 'age', '0.093', 'Younger patients have higher risk — consistent with more aggressive early-onset disease.'],
        ['6', 'bmi', '0.081', 'Despite 77% missingness, BMI contributes. Higher BMI pushes toward uncontrolled.'],
        ['7', 'a1c_variability', '0.078', 'Unstable A1c (high standard deviation) increases risk.'],
        ['8', 'a1c_above_9', '0.057', 'The binary threshold for severe uncontrolled diabetes.'],
        ['9', 'is_male', '0.055', 'Being male modestly increases predicted risk.'],
        ['10', 'sulfonylurea orders', '0.047', 'Sulfonylurea use pushes toward uncontrolled — confounded (prescribed to sicker patients).'],
    ]
)

add_image('plots/step13_shap_bar.png', 5.0)
doc.add_paragraph('Figure 19: Mean absolute SHAP values — the simplified importance ranking.')

bold_para('The clinical story:')
doc.add_paragraph(
    'The model\'s decision-making mirrors clinical reasoning. It first evaluates A1c history — '
    'the mean, the worst episode, the most recent reading, and whether things are getting better '
    'or worse. It then considers treatment response (medications despite high A1c), demographics '
    '(age, sex), and social context (ADI, insurance). This interpretability is a key strength for '
    'clinical adoption.'
)

# --- Step 14 ---
doc.add_heading('5.2 Step 14: Subgroup Fairness Analysis', level=2)

doc.add_paragraph(
    'A clinical prediction model must work equitably across populations. If a model performs well '
    'for White patients but poorly for Black patients, deploying it would worsen health disparities '
    'rather than reduce them. We tested model performance across race, gender, age, insurance, and '
    'neighborhood deprivation (ADI).'
)

bold_para('Results by race:')
add_table(
    ['Race', 'N (test)', 'AUC-ROC', 'Recall', 'Assessment'],
    [
        ['Asian', '245', '0.880', '86.7%', 'Best performance'],
        ['Black or African American', '3,395', '0.860', '84.8%', 'Strong — catches MORE than White'],
        ['White', '8,672', '0.847', '78.5%', 'Solid baseline'],
    ]
)

bold_para('Results by gender:')
add_table(
    ['Gender', 'N (test)', 'AUC-ROC', 'Recall'],
    [
        ['Female', '6,498', '0.872', '82.7%'],
        ['Male', '5,986', '0.829', '78.8%'],
    ]
)

bold_para('Results by insurance:')
doc.add_paragraph(
    'Medicaid patients — the most vulnerable population — had the highest recall: 91–93%. The '
    'model is especially sensitive for this group. Medicare patients had the highest AUC (0.871–0.895).'
)

bold_para('Results by ADI (neighborhood deprivation):')
doc.add_paragraph(
    'AUC ranged from 0.839 to 0.881 across ADI quartiles. The least deprived neighborhoods had '
    'slightly better AUC, but the gap is small (0.042). The model does not meaningfully disadvantage '
    'patients from deprived areas.'
)

add_image('plots/step14_subgroup_analysis.png', 5.5)
doc.add_paragraph('Figure 20: AUC and recall across race, gender, age, insurance, and ADI subgroups.')

bold_para('Fairness assessment:')
doc.add_paragraph(
    'The model does not disadvantage minority or low-income subgroups. If anything, it has higher '
    'sensitivity for Black patients and Medicaid patients — populations that face the greatest '
    'diabetes disparities. AUC ranges from 0.83 to 0.90 across all tested subgroups. For '
    'real-world deployment, this means the tool could help close equity gaps rather than widen them.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6: SUBMISSION AND PRESENTATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Submission and Presentation', level=1)

doc.add_heading('Submission', level=2)
doc.add_paragraph(
    'We generated two files. The primary submission file (submission.csv) matches the target file '
    'format exactly: patient ID as the index, one column "a1c 2025 Uncontrolled" with True/False '
    'values, 62,425 rows. A detailed file (submission_detailed.csv) includes the predicted '
    'probability for each patient, the predicted class, and the actual outcome for reference.'
)

doc.add_paragraph(
    'At threshold 0.5, the model flags 19,694 patients (31.5%) as uncontrolled while the actual '
    'rate is 10.6%. This is by design — the model prioritizes catching at-risk patients over '
    'minimizing false alarms, which is the appropriate tradeoff for a screening tool.'
)

doc.add_heading('Presentation Strategy', level=2)
doc.add_paragraph(
    'The presentation follows a 13-slide arc: clinical problem → data → key insight (A1c risk '
    'gradient) → feature engineering → model comparison → final performance → interpretability '
    '(SHAP) → equity → clinical application → limitations → summary. Seven publication-ready '
    'figures were created with the "pres_" prefix in the plots folder.'
)

doc.add_paragraph(
    'Key messages: (1) A1c history is the dominant predictor but insufficient alone. '
    '(2) Engineered features like treatment_resistant capture clinically meaningful interactions. '
    '(3) The model achieves AUC 0.852 with 81% recall and 97% NPV. '
    '(4) Performance is equitable across demographic subgroups.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7: KEY DECISIONS LOG
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Key Decisions Log', level=1)

doc.add_paragraph(
    'This table documents every major decision we made during the project, the reasoning behind it, '
    'and where in the pipeline it occurred.'
)

add_table(
    ['Step', 'Decision', 'Reasoning'],
    [
        ['Step 4', 'Drop the leakage column (a1c 2025 collection date)', '100% of patients missing this value are Controlled. The model could achieve near-perfect accuracy by checking this column — but that would not generalize to new patients where this information is not available.'],
        ['Step 5', 'Drop ADI state rank, keep national rank', 'The two are 96% correlated (r = 0.964). Keeping both adds no information and introduces multicollinearity. National rank is more standardized.'],
        ['Step 5', 'Aggregate raw A1c columns into summary features', 'The 5 raw A1c columns are correlated with each other (r = 0.52–0.64) and 58–97% missing. Summary features (mean, max, change, variability) are more complete and clinically interpretable.'],
        ['Step 6', 'Simplify insurance from 12 to 4 categories', 'Granular plan names (e.g., "Blue Cross Managed Care") add noise. Broader categories (Medicare, Medicaid, Managed Care, Other) capture what matters: public vs private vs managed coverage type.'],
        ['Step 7', 'Create treatment_resistant and undertreated features', 'Raw medication counts are confounded by disease severity. Combining A1c level with medication count creates a clinically meaningful signal: is the patient failing despite treatment, or untreated despite need?'],
        ['Step 8', 'Use stratified 80/20 split with random_state=42', 'Stratification preserves the 10.6% class balance in both sets. Fixed random state ensures reproducibility. 80/20 is standard for datasets of this size.'],
        ['Step 9', 'Use class_weight="balanced" for logistic regression', 'Without class weighting, the model predicts "controlled" for everyone (89.4% accuracy, 0% recall). Balanced weighting tells the model that missing an uncontrolled patient is 8× worse than a false alarm.'],
        ['Step 10', 'Choose XGBoost for tuning over Random Forest', 'XGBoost typically outperforms Random Forest once tuned. It supports native NaN handling (no imputation needed) and has more tunable hyperparameters.'],
        ['Step 11', 'Use scale_pos_weight = 8.41', 'This is the ratio of controlled to uncontrolled patients (44,634 / 5,306). It serves the same purpose as class_weight="balanced" — penalizes false negatives proportionally to the imbalance.'],
        ['Step 11', 'learning_rate = 0.01 with 800 trees', 'Slower learning rate + more trees = better generalization. Each tree makes a small correction, reducing overfitting. This is the most impactful tuning choice.'],
        ['Step 11', 'subsample = 0.7, colsample_bytree = 0.5', 'Randomness in training prevents overfitting. Each tree sees 70% of patients and 50% of features, forcing the model to learn robust patterns rather than memorizing.'],
        ['Step 12', 'Use threshold = 0.5 as default', 'Standard threshold for binary classification. Can be adjusted based on clinical context (lower for screening, higher for intervention). The threshold table provides options.'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8: GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Glossary', level=1)

doc.add_paragraph(
    'Every machine learning term used in this document, defined in plain English with a clinical '
    'analogy where applicable.'
)

glossary = [
    ['AUC-ROC (Area Under the Receiver Operating Characteristic Curve)',
     'A single number (0 to 1) that measures how well the model separates the two classes. '
     'An AUC of 0.85 means: pick a random uncontrolled and a random controlled patient — 85% of '
     'the time, the model gives the uncontrolled patient a higher risk score. '
     'Clinical analogy: like the sensitivity/specificity tradeoff of a diagnostic test, but '
     'summarized across all possible thresholds at once. 0.5 = random, 0.7–0.8 = acceptable, '
     '0.8–0.9 = good, 0.9+ = excellent.'],

    ['Sensitivity (Recall)',
     'Of all patients who ARE uncontrolled, what percentage does the model correctly identify? '
     'Our model\'s sensitivity of 80.6% means it catches 4 out of 5 at-risk patients. '
     'Clinical analogy: like the sensitivity of a rapid strep test — does it find the cases?'],

    ['Specificity',
     'Of all patients who ARE controlled, what percentage does the model correctly clear? '
     'Our specificity of 74.2% means it correctly identifies 3 out of 4 healthy patients. '
     'Clinical analogy: like the specificity of a pregnancy test — does a negative mean you\'re really not pregnant?'],

    ['Precision (Positive Predictive Value / PPV)',
     'Of all patients the model FLAGS as uncontrolled, what percentage actually are? '
     'Our precision of 27% means 1 in 4 flagged patients is truly uncontrolled. '
     'This depends heavily on the base rate (prevalence). With only 10.6% truly uncontrolled, '
     'even a good model generates many false alarms — the same reason mammography has low PPV.'],

    ['NPV (Negative Predictive Value)',
     'Of all patients the model CLEARS as low-risk, what percentage truly stay controlled? '
     'Our NPV of 97% means: if the model says you\'re safe, there\'s a 97% chance it\'s right. '
     'This is the model\'s strongest clinical metric — its "rule-out" power.'],

    ['F1 Score',
     'The harmonic mean of precision and recall — a single number that balances both. '
     'Useful when you want a simple summary of the precision/recall tradeoff. '
     'Ranges from 0 to 1.'],

    ['PR-AUC (Precision-Recall Area Under Curve)',
     'Like AUC-ROC but focuses on the positive class (uncontrolled patients). '
     'More informative than AUC-ROC for imbalanced datasets because it measures how well '
     'the model finds the rare positive cases without too many false alarms.'],

    ['Brier Score',
     'Measures calibration — how accurate the probability estimates are. '
     'If the model says "70% chance of uncontrolled," how often is that actually true? '
     'Ranges from 0 (perfect) to 1 (worst). Our score of 0.159 indicates reasonable calibration.'],

    ['Overfitting',
     'When a model memorizes the training data instead of learning general patterns. '
     'An overfit model performs well on training data but poorly on new patients. '
     'Clinical analogy: a medical student who memorizes practice questions word-for-word '
     'but cannot handle rephrased questions on the exam.'],

    ['Train-Test Split',
     'Dividing data into a training set (used to build the model) and a test set (used to '
     'evaluate it on unseen data). The test set is like a sealed exam — if you peek during '
     'studying, your score does not reflect real knowledge.'],

    ['Cross-Validation (CV)',
     'A technique where the training data is split into K folds. The model is trained on K-1 '
     'folds and evaluated on the remaining fold, rotating K times. This gives a more reliable '
     'performance estimate than a single split. We used 5-fold CV during hyperparameter tuning.'],

    ['Feature Engineering',
     'Creating new variables from raw data that better capture the underlying signal. '
     'Example: instead of using 5 raw A1c values (most missing), we computed a1c_mean, a1c_max, '
     'and a1c_change — features that are more complete and clinically meaningful.'],

    ['One-Hot Encoding',
     'Converting a categorical variable (like race) into multiple binary (0/1) columns — one '
     'per category. Patient is Asian → race_Asian = 1, all other race columns = 0. '
     'Necessary because most ML models require numeric input.'],

    ['XGBoost (Extreme Gradient Boosting)',
     'An ensemble method that builds many small decision trees sequentially. Each tree corrects '
     'the mistakes of the previous ones. Known for high accuracy on tabular data. '
     'Clinical analogy: like a tumor board where each specialist adds a correction to the '
     'diagnosis, building on the previous specialists\' assessments.'],

    ['Random Forest',
     'An ensemble of many decision trees, each trained on a random subset of data and features. '
     'The final prediction is the majority vote. Simpler than XGBoost but often nearly as good. '
     'Clinical analogy: getting opinions from many independent doctors and going with the consensus.'],

    ['Logistic Regression',
     'The simplest classification model. Finds a linear combination of features that predicts '
     'the probability of the positive class. Highly interpretable but limited to linear relationships. '
     'Clinical analogy: a risk score calculator where each factor adds or subtracts points.'],

    ['SHAP Values (SHapley Additive exPlanations)',
     'A method from game theory that explains each individual prediction. For each patient, SHAP '
     'shows exactly how much each feature pushed the prediction toward "uncontrolled" or "controlled." '
     'Clinical analogy: like breaking down a diagnosis into contributing factors — "your high A1c '
     'contributes +3 points of risk, your young age contributes +1, your medication contributes −0.5."'],

    ['Hyperparameter Tuning',
     'The process of finding the best model settings (hyperparameters) by systematically trying '
     'different combinations and evaluating each with cross-validation. '
     'Clinical analogy: titrating a drug dose — trying different levels to find the sweet spot '
     'between efficacy and side effects.'],

    ['Class Imbalance',
     'When one class is much more common than the other. In our data, 89.4% are controlled and '
     'only 10.6% are uncontrolled. Without special handling, models learn to always predict the '
     'majority class. We addressed this with class_weight="balanced" and scale_pos_weight.'],

    ['scale_pos_weight',
     'An XGBoost parameter that sets the cost ratio between missing a positive (uncontrolled) '
     'case and raising a false alarm. Our value of 8.41 means: missing an uncontrolled patient '
     'costs 8.41× more than incorrectly flagging a controlled patient. This matches the ~1:8 '
     'class imbalance ratio.'],
]

for term, definition in glossary:
    p = doc.add_paragraph()
    run = p.add_run(term)
    run.bold = True
    p.add_run('\n' + definition)
    p.paragraph_format.space_after = Pt(8)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save('datathon_notebook.docx')
print('Saved: datathon_notebook.docx')
