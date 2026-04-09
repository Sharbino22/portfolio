"""Build case_study.docx with findings and embedded figures."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

FIGURES = Path('figures')
OUT = Path('docs/case_study.docx')

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0F, 0x0F, 0x14)


def add_figure(path, caption, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    cap.runs[0].font.italic = True


# ── Title ──
title = doc.add_heading('Medicare Rate and Utilization Analysis', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('A Case Study in Payment Variation Across Providers, Procedures, and Geographies')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Sridharan Gopalsamy Ramaswamy | MPH/MBA, Washington University in St. Louis | April 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# ── Executive Summary ──
doc.add_heading('Executive Summary', level=2)
doc.add_paragraph(
    'This analysis examines Medicare payment patterns across 3,015 inpatient hospitals '
    'and 1.15 million physician providers using CMS public use files for Calendar Year 2022. '
    'Three findings stand out for anyone working in payer strategy, provider contracting, '
    'or health economics:'
)
doc.add_paragraph(
    'Hospitals charge 5.4x what Medicare pays (median), but this ratio varies from '
    '1.2x in Maryland to 10.7x in Nevada, a 9.5x spread that reflects state-level '
    'regulatory and market dynamics.',
    style='List Bullet'
)
doc.add_paragraph(
    'Septicemia dominates inpatient volume, accounting for 550K discharges (11% of all '
    'Medicare inpatient stays), while heart transplants lead in cost at ~$295K per case.',
    style='List Bullet'
)
doc.add_paragraph(
    'Physician markup ratios vary sharply by specialty, with certain specialties billing '
    '4-5x the Medicare allowed amount, signaling where payer-provider rate negotiations '
    'have the most room to move.',
    style='List Bullet'
)

# ── Data Overview ──
doc.add_heading('Data Overview', level=2)
table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
headers = ['Metric', 'Inpatient', 'Physician']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('Records', '145,742', '9,755,020'),
    ('Providers', '3,015 hospitals', '1,148,808 NPIs'),
    ('Procedures', '533 DRGs', '6,326 HCPCS'),
    ('States', '51', '61'),
    ('Median markup', '5.37x', '2.79x'),
]
for r, (metric, inp_val, phy_val) in enumerate(data, 1):
    table.rows[r].cells[0].text = metric
    table.rows[r].cells[1].text = inp_val
    table.rows[r].cells[2].text = phy_val

doc.add_paragraph()

# ── Section 1: Volume ──
doc.add_heading('1. Where the Volume Is: Top DRGs', level=2)
doc.add_paragraph(
    'Septicemia without MV >96 hours leads all DRGs with 550,306 discharges, followed '
    'by heart failure (324,750) and respiratory infections (275,104). The top 5 DRGs '
    'account for a disproportionate share of Medicare inpatient volume.'
)
add_figure(FIGURES / 'top15_drg_volume.png', 'Figure 1. Top 15 DRGs by Discharge Volume (2022)')
doc.add_paragraph(
    'Strategic implication: Any value-based contract that does not explicitly address '
    'septicemia, heart failure, and respiratory infections is leaving the highest-volume '
    'conditions unmanaged. Bundled payment and readmission reduction programs should '
    'prioritize these pathways.'
)

# ── Section 2: Cost ──
doc.add_heading('2. Where the Money Goes: Costliest DRGs', level=2)
doc.add_paragraph(
    'Heart transplants and ECMO cases average $295K and $178K per case respectively. '
    'For commercial payers benchmarking against Medicare, these are the DRGs where a '
    '10% rate difference translates to $20K-$30K per case.'
)
add_figure(FIGURES / 'top15_drg_cost.png', 'Figure 2. Top 15 Costliest DRGs by Avg Medicare Payment (2022)')

# ── Section 3: Markup ──
doc.add_heading('3. The Markup Gap: Charges vs. Medicare Payment', level=2)
doc.add_paragraph(
    'Nationally, the median hospital charges 5.4x what Medicare ultimately pays. '
    'The scatter plot reveals a nonlinear relationship: for higher-cost DRGs, the '
    'markup ratio compresses, while lower-cost procedures show the widest spread. '
    'The distribution is right-skewed, with a minority of hospitals charging 8-12x Medicare rates.'
)
add_figure(FIGURES / 'inpatient_charges_vs_payment.png', 'Figure 3. Submitted Charges vs. Medicare Payment (Inpatient)')
add_figure(FIGURES / 'inpatient_markup_distribution.png', 'Figure 4. Distribution of Inpatient Markup Ratios')
doc.add_paragraph(
    'Strategic implication: Chargemaster-based contracts systematically overpay relative to '
    'Medicare. Percent-of-Medicare benchmarks are more defensible for payer contracting.'
)

# ── Section 4: Geography ──
doc.add_heading('4. Geographic Rate Variation', level=2)
doc.add_paragraph(
    'The 9.5x spread in median markup between Maryland (1.2x) and Nevada (10.7x) '
    'is not random. Maryland operates an all-payer rate-setting system, compressing the '
    'charge-to-payment gap by design. States without rate regulation show far higher markups.'
)
add_figure(FIGURES / 'state_markup_variation.png', 'Figure 5. Inpatient Markup Ratio by State (Median with IQR)')
add_figure(FIGURES / 'state_volume_payment.png', 'Figure 6. Discharge Volume and Avg Payment by State')
add_figure(FIGURES / 'state_payment_boxplot.png', 'Figure 7. Medicare Payment Distribution in Top 10 States by Volume')
doc.add_paragraph(
    'Strategic implication: Multi-state payers need state-specific rate playbooks. '
    'States like California, Texas, and Florida have high volume and above-median markups, '
    'making them priority markets for rate renegotiation.'
)

# ── Section 5: Physician Volume + Cost ──
doc.add_heading('5. Physician Services: Volume and Cost Patterns', level=2)
doc.add_paragraph(
    'Clinical laboratory and hematology-oncology dominate service volume, while '
    'ambulatory surgical centers and cardiac surgery lead in per-service cost. '
    'The volume-cost mismatch by specialty is a signal for payers: high-volume specialties '
    'drive spend through utilization, high-cost specialties drive it through unit price.'
)
add_figure(FIGURES / 'top15_specialty_volume.png', 'Figure 8. Top 15 Specialties by Service Volume')
add_figure(FIGURES / 'top15_specialty_cost.png', 'Figure 9. Top 15 Costliest Specialties by Avg Payment')

# ── Section 6: Physician Rate Variation ──
doc.add_heading('6. Physician Markup and Rate Variation', level=2)
doc.add_paragraph(
    'Physician-level analysis shows a median markup of 2.8x (submitted charge / Medicare allowed). '
    'Certain specialties bill 4-5x the allowed amount. The HCPCS-level coefficient of variation '
    'chart identifies specific procedures where payment varies most across providers.'
)
add_figure(FIGURES / 'physician_charges_vs_allowed.png', 'Figure 10. Submitted Charges vs. Medicare Allowed (Physician)')
add_figure(FIGURES / 'specialty_markup_ratio.png', 'Figure 11. Markup Ratio by Specialty (Top 20)')
add_figure(FIGURES / 'hcpcs_payment_variation.png', 'Figure 12. Highest Payment Variation Across Procedures (HCPCS)')
add_figure(FIGURES / 'state_physician_rate_spread.png', 'Figure 13. Physician Rate Spread by State')

# ── Strategic Implications ──
doc.add_heading('Strategic Implications', level=2)
impl_table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
impl_table.rows[0].cells[0].text = 'Finding'
impl_table.rows[0].cells[1].text = 'Strategic Implication'
impl_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
impl_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

implications = [
    ('Septicemia = 11% of inpatient volume',
     'Bundled payment and readmission reduction programs should prioritize sepsis pathways'),
    ('Median inpatient markup = 5.4x',
     'Chargemaster-based contracts systematically overpay; percent-of-Medicare benchmarks are more defensible'),
    ('MD vs. NV markup spread = 9.5x',
     'State regulatory environment is a first-order variable; multi-state payers need state-specific playbooks'),
    ('High-variation HCPCS codes identified',
     'Procedure-level rate benchmarking can surface outlier providers billing 3-5x peers'),
    ('Volume-cost specialty mismatch',
     'Utilization management for high-volume specialties; rate negotiation for high-cost specialties'),
]
for r, (finding, implication) in enumerate(implications, 1):
    impl_table.rows[r].cells[0].text = finding
    impl_table.rows[r].cells[1].text = implication

doc.add_paragraph()

# ── Methods ──
doc.add_heading('Methods', level=2)
doc.add_paragraph(
    'Data: CMS Medicare Provider Utilization and Payment PUF, Calendar Year 2022. '
    'Inpatient file: 145,742 provider-DRG records across 3,015 hospitals and 533 DRGs. '
    'Physician file: 9.76M provider-HCPCS records across 1.15M NPIs and 6,326 procedure codes. '
    'Cleaning: deduplication on natural keys, encoding normalization (Latin-1), US-only filter '
    'for physician file. Derived variable: markup ratio = submitted charges / Medicare payment '
    '(inpatient) or submitted charges / Medicare allowed (physician). Tools: Python, pandas, matplotlib.'
)

# ── Footer ──
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Analysis by Sridharan Gopalsamy Ramaswamy | sridharanshri.com')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save(str(OUT))
print(f'Saved: {OUT} ({OUT.stat().st_size / 1e6:.1f} MB)')
