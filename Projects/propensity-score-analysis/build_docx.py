"""Build case_study.docx with findings and embedded figures."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

FIGURES = Path('figures')
OUT = Path('docs/case_study.docx')

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0F, 0x0F, 0x14)


def add_figure(path, caption, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    cap.runs[0].font.italic = True


# Title
title = doc.add_heading('Propensity Score Analysis of Antihypertensive Treatment\nand Cardiovascular Mortality', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('A Comparative Effectiveness Study Using NHANES Linked Mortality Data')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Sridharan Gopalsamy Ramaswamy | MPH/MBA, Washington University in St. Louis | April 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', level=2)
doc.add_paragraph(
    'This study estimates the effect of antihypertensive treatment on cardiovascular mortality '
    'in 18,129 hypertensive US adults using propensity score methods applied to NHANES 1999-2018 '
    'with linked mortality follow-up through 2019.'
)
doc.add_paragraph(
    'Confounding is massive. The unadjusted HR of 2.86 suggests treatment increases CV death, '
    'but this is an artifact of confounding by indication. After PS matching, the HR attenuates to 1.24.',
    style='List Bullet'
)
doc.add_paragraph(
    'Method choice changes the conclusion. PS matching finds HR 1.24 (p=0.027), IPTW finds no effect '
    '(HR 1.02, p=0.81), and the doubly robust estimate is HR 1.15 (p=0.09).',
    style='List Bullet'
)
doc.add_paragraph(
    'The PH assumption is borderline violated (p=0.035), suggesting the treatment effect '
    'varies over time, consistent with delayed benefit of BP control.',
    style='List Bullet'
)

# Data Overview
doc.add_heading('Data Overview', level=2)
table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
headers = ['Metric', 'Full Cohort', 'PS-Matched']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
data = [
    ('N', '18,129', '6,962'),
    ('Treated', '13,649 (75.3%)', '3,481 (50.0%)'),
    ('CV Deaths', '1,641 (9.1%)', '577 (8.3%)'),
    ('Median follow-up', '7.8 years', '9.3 years'),
]
for r, (m, v1, v2) in enumerate(data, 1):
    table.rows[r].cells[0].text = m
    table.rows[r].cells[1].text = v1
    table.rows[r].cells[2].text = v2
doc.add_paragraph()

# Section 1: PS Distribution
doc.add_heading('1. Propensity Score Distribution', level=2)
doc.add_paragraph(
    'The propensity score model (logistic regression on 12 covariates) achieves 81% accuracy. '
    'The score distributions show substantial overlap with clear separation, confirming systematic '
    'confounding by indication.'
)
add_figure(FIGURES / 'ps_distribution.png', 'Figure 1. Propensity Score Distribution by Treatment Group')

# Section 2: Balance
doc.add_heading('2. Covariate Balance', level=2)
doc.add_paragraph(
    'Both PS matching and IPTW reduce standardized mean differences substantially. '
    'Age, the strongest confounder (SMD ~0.70 unadjusted), drops below 0.1 after matching.'
)
add_figure(FIGURES / 'love_plot.png', 'Figure 2. Love Plot: Covariate Balance Before and After Adjustment')

# Section 3: KM Curves
doc.add_heading('3. Survival Analysis', level=2)
doc.add_paragraph(
    'The unadjusted KM shows dramatic separation (treated patients dying faster because they are sicker). '
    'After matching, the curves converge substantially.'
)
add_figure(FIGURES / 'km_curves.png', 'Figure 3. Kaplan-Meier Curves: Unadjusted and PS-Matched')

# Section 4: HR Comparison
doc.add_heading('4. Treatment Effect Estimates', level=2)
doc.add_paragraph(
    'Four Cox PH models estimate the treatment hazard ratio under increasingly rigorous confounding control.'
)
add_figure(FIGURES / 'hr_forest_plot.png', 'Figure 4. Hazard Ratio Forest Plot Across PS Methods')

hr_table = doc.add_table(rows=5, cols=5, style='Light Grid Accent 1')
hr_headers = ['Model', 'HR', '95% CI', 'p', 'N']
for i, h in enumerate(hr_headers):
    hr_table.rows[0].cells[i].text = h
    hr_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
hr_data = [
    ('Unadjusted', '2.86', '2.47-3.32', '<0.001', '18,129'),
    ('PS-Matched', '1.24', '1.03-1.50', '0.027', '6,962'),
    ('IPTW', '1.02', '0.87-1.21', '0.808', '18,129'),
    ('Doubly Robust', '1.15', '0.98-1.36', '0.091', '18,129'),
]
for r, vals in enumerate(hr_data, 1):
    for c, v in enumerate(vals):
        hr_table.rows[r].cells[c].text = v
doc.add_paragraph()

doc.add_paragraph(
    'The IPTW estimate (HR 1.02) is the most fully adjusted and suggests no net harm or benefit. '
    'However, this should not be interpreted as treatment futility. Residual confounding, '
    'the healthy-user effect, and limitations of self-reported treatment data are likely explanations.'
)

# Section 5: PH Test
doc.add_heading('5. Proportional Hazards Assessment', level=2)
doc.add_paragraph(
    'The HR-over-time plot shows the treatment effect is not constant: higher in early follow-up, '
    'attenuating over time. Formal PH test: p=0.035 for treatment.'
)
add_figure(FIGURES / 'ph_test_plot.png', 'Figure 5. Treatment HR Over Time')

# Section 6: Subgroup
doc.add_heading('6. Subgroup Analysis', level=2)
doc.add_paragraph(
    'The subgroup forest plot tests whether the treatment effect varies by age, sex, diabetes, '
    'and smoking status in the PS-matched cohort.'
)
add_figure(FIGURES / 'subgroup_forest_plot.png', 'Figure 6. Subgroup Forest Plot (PS-Matched)')

# Implications
doc.add_heading('Implications for HEOR and Consulting', level=2)
impl_table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
impl_table.rows[0].cells[0].text = 'Finding'
impl_table.rows[0].cells[1].text = 'Implication'
impl_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
impl_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
implications = [
    ('Unadjusted HR reverses after PS adjustment',
     'Observational analyses require rigorous confounding control for formulary/coverage decisions'),
    ('HR varies 1.02-1.24 across methods',
     'Sensitivity analysis across PS approaches is non-negotiable for HEOR submissions'),
    ('PH assumption borderline violated',
     'Time-varying treatment effects should be explored; standard Cox may mask delayed benefit'),
    ('3,481 matched pairs from 18,129',
     'IPTW preserves the full sample and may be preferable when overlap is limited'),
    ('Self-reported treatment is a limitation',
     'Claims-based identification (NDC codes, pharmacy fills) would strengthen real-world analyses'),
]
for r, (f, i) in enumerate(implications, 1):
    impl_table.rows[r].cells[0].text = f
    impl_table.rows[r].cells[1].text = i
doc.add_paragraph()

# Methods
doc.add_heading('Methods', level=2)
doc.add_paragraph(
    'Data: NHANES 1999-2018 (10 cycles) linked to NCHS mortality (follow-up through Dec 2019). '
    'Cohort: 18,129 adults 20+ with self-reported hypertension. '
    'Treatment: antihypertensive medication use (BPQ050A). '
    'Outcome: CV mortality (UCOD heart disease or cerebrovascular disease). '
    'PS model: logistic regression on 12 covariates. '
    'Matching: 1:1 nearest-neighbor, caliper = 0.2 SD of logit PS. '
    'IPTW: ATE weights, trimmed at 99th percentile. '
    'Tools: Python, lifelines, scikit-learn, pandas, matplotlib.'
)

# Footer
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Analysis by Sridharan Gopalsamy Ramaswamy | sridharanshri.com')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save(str(OUT))
print(f'Saved: {OUT} ({OUT.stat().st_size / 1e6:.1f} MB)')
